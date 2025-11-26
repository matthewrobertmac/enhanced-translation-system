# src/utils/files.py
import io
from datetime import datetime

# Optional imports with protection
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def read_uploaded_file(uploaded_file) -> str:
    """
    Read content from uploaded file (supports txt, md, docx).
    """
    if uploaded_file is None:
        return ""
        
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension in ('txt', 'md'):
            return uploaded_file.read().decode('utf-8')
            
        elif file_extension == 'docx':
            if not DOCX_AVAILABLE:
                return "[Error: python-docx not installed. Install with: pip install python-docx]"
            doc = Document(uploaded_file)
            return '\n\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            
        else:
            return f"[Error: Unsupported file type: {file_extension}]"
            
    except Exception as e:
        return f"[Error reading file: {str(e)}]"

def create_docx_file(text: str, title: str = "Translation") -> io.BytesIO:
    """
    Create a formatted Word document stream.
    """
    if not DOCX_AVAILABLE:
        return None
    
    doc = Document()
    title_paragraph = doc.add_heading(title, level=1)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    timestamp = doc.add_paragraph()
    timestamp.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    for para in text.split('\n\n'):
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def create_markdown_file(text: str, title: str = "Translation") -> str:
    """
    Create a formatted Markdown string.
    """
    return f"""# {title}

**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

---

{text}
"""