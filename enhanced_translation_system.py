"""
Enhanced Multi-Agent Translation Workflow with LangGraph and LangSmith
A sophisticated translation system with cultural adaptation, literary editing, comprehensive monitoring, and visuals.

Features:
- INTELLIGENT PLANNING AGENT - dynamically selects required agents
- 7 specialized translation agents with distinct roles (including BERTScore validator)
- **SMART SEMANTIC CACHING** - 5-10x speedup on similar content
- **CONFIDENCE SCORES** - Multi-metric translation quality assessment
- **DIFF VISUALIZATION** - Visual comparison between agent versions
- **ALTERNATIVE TRANSLATIONS** - Generate and compare multiple variants
- Support for OpenAI and Anthropic models
- Optional LangSmith tracing and monitoring
- Comprehensive agent feedback system
- File upload support (txt, docx, md)
- Multiple export formats (txt, docx, md)
- Critical passage flagging and review
- Safe same-language (e.g., English→English) refinement mode
- BERTScore validation with iterative refinement
- Visualizations: word counts, sentence-length histograms, readability, issue counts, BERTScore bars
- Word clouds: Source, Final, and Difference (words added)
- Entity tracking and network visualization
"""
import streamlit as st
from typing import TypedDict, Annotated, List, Dict, Optional, Literal
from langgraph.graph import StateGraph, END
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langchain_core.language_models import BaseChatModel
import operator
from datetime import datetime
import json
import os
import io
import traceback
import re
import collections
import difflib
import asyncio
import hashlib
import pickle

# === New: viz imports ===
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt

# === New: wordcloud (optional) ===
try:
    from wordcloud import WordCloud, STOPWORDS
    WORDCLOUD_AVAILABLE = True
except Exception:
    WORDCLOUD_AVAILABLE = False

# === Entity tracking imports ===
try:
    import networkx as nx
    NETWORKX_AVAILABLE = True
except ImportError:
    NETWORKX_AVAILABLE = False

try:
    import plotly.express as px
    import plotly.graph_objects as go
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False

# === Sentence embeddings for caching ===
try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

# ===== Language Guardrail =====
LANGUAGE_GUARDRAIL = (
    "STRICT LANGUAGE GUARDRAIL:\n"
    "- Reply ONLY in the target language specified.\n"
    "- Do NOT include any words or phrases in other languages.\n"
    "- Any notes, bullets, or headings must also be in the target language.\n"
)

# =====================
# CONFIDENCE SCORER
# =====================
class ConfidenceScorer:
    """Score translation confidence using multiple signals"""
    
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
    
    async def score_translation(
        self, 
        source: str, 
        translation: str, 
        context: dict
    ) -> Dict[str, float]:
        """Generate confidence scores"""
        
        scores = {}
        
        # 1. Semantic similarity (if same-language or embeddings available)
        try:
            from bert_score import score as bert_score_fn
            if languages_equivalent(context.get('source_lang', ''), context.get('target_lang', '')):
                P, R, F1 = bert_score_fn([translation], [source], lang="en", rescale_with_baseline=True)
                scores['semantic_fidelity'] = float(F1.mean().item())
        except:
            scores['semantic_fidelity'] = 0.75  # Default reasonable score
        
        # 2. Length ratio (should be reasonable)
        src_len = len(source.split())
        tgt_len = len(translation.split())
        ratio = tgt_len / src_len if src_len > 0 else 0
        # Ideal ratio is 0.8-1.2
        scores['length_ratio'] = max(0, 1.0 - min(abs(ratio - 1.0), 0.5) * 2)
        
        # 3. Fluency (use language model perplexity proxy)
        prompt = f"""Rate the fluency/naturalness of this {context.get('target_lang', 'target')} text on a scale of 0.0-1.0.

Text: {translation[:500]}

Respond with ONLY a number between 0.0 and 1.0, no explanation."""
        
        try:
            response = await self.llm.ainvoke([HumanMessage(content=prompt)])
            fluency_str = response.content.strip()
            # Extract first number found
            import re
            numbers = re.findall(r'0\.\d+|1\.0|0|1', fluency_str)
            if numbers:
                fluency = float(numbers[0])
                scores['fluency'] = min(max(fluency, 0.0), 1.0)
            else:
                scores['fluency'] = 0.75
        except:
            scores['fluency'] = 0.75  # Default to reasonable score
        
        # 4. Terminology consistency
        entities = context.get('entities', [])
        if entities:
            preserved = sum(1 for e in entities if e['name'].lower() in translation.lower())
            total = len(entities)
            scores['terminology'] = preserved / total if total > 0 else 1.0
        else:
            scores['terminology'] = 0.85  # Default when no entities
        
        # 5. Overall confidence (weighted average)
        weights = {
            'semantic_fidelity': 0.35,
            'length_ratio': 0.15,
            'fluency': 0.35,
            'terminology': 0.15
        }
        
        overall = sum(scores.get(k, 0.5) * v for k, v in weights.items())
        scores['overall'] = overall
        
        return scores

# =====================
# DIFF VISUALIZATION
# =====================
def create_diff_visualization(text1: str, text2: str, title1: str, title2: str) -> str:
    """Visual diff with highlighting"""
    
    # Split into words for better granularity
    words1 = text1.split()
    words2 = text2.split()
    
    d = difflib.Differ()
    diff = list(d.compare(words1, words2))
    
    html_output = f"""
    <div style='background: #f8f9fa; padding: 20px; border-radius: 8px; margin: 10px 0;'>
        <div style='display: grid; grid-template-columns: 1fr 1fr; gap: 20px;'>
            <div>
                <h4 style='color: #dc3545; margin-bottom: 10px;'>← {title1}</h4>
                <div style='background: white; padding: 15px; border-radius: 5px; border-left: 4px solid #dc3545; line-height: 1.8;'>
    """
    
    # Build left side (removals and unchanged)
    for word in diff:
        if word.startswith('- '):
            html_output += f"<span style='background: #ffcdd2; padding: 2px 4px; margin: 1px; text-decoration: line-through;'>{word[2:]}</span> "
        elif word.startswith('  '):
            html_output += f"<span style='color: #666;'>{word[2:]}</span> "
    
    html_output += """
                </div>
            </div>
            <div>
                <h4 style='color: #28a745; margin-bottom: 10px;'>{} →</h4>
                <div style='background: white; padding: 15px; border-radius: 5px; border-left: 4px solid #28a745; line-height: 1.8;'>
    """.format(title2)
    
    # Build right side (additions and unchanged)
    for word in diff:
        if word.startswith('+ '):
            html_output += f"<span style='background: #c8e6c9; padding: 2px 4px; margin: 1px; font-weight: 500;'>{word[2:]}</span> "
        elif word.startswith('  '):
            html_output += f"<span style='color: #666;'>{word[2:]}</span> "
    
    html_output += """
                </div>
            </div>
        </div>
    </div>
    """
    
    return html_output

def calculate_change_rate(text1: str, text2: str) -> float:
    """Calculate percentage of text that changed"""
    similarity = difflib.SequenceMatcher(None, text1, text2).ratio()
    return (1 - similarity) * 100

# =====================
# ALTERNATIVE TRANSLATION GENERATOR
# =====================
class AlternativeTranslationGenerator:
    """Generate alternative translation variants"""
    
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
    
    async def generate_alternatives(
        self,
        state: "TranslationState",  # Use string for forward reference
        num_alternatives: int = 3
    ) -> List[Dict]:
        """Generate alternative translations with different strategies"""
        
        alternatives = []
        
        # Strategy 1: More conservative (lower temperature)
        alternatives.append(await self._generate_variant(
            state,
            temperature=0.2,
            strategy="conservative",
            description="More literal and conservative"
        ))
        
        # Strategy 2: Balanced (medium temperature)
        alternatives.append(await self._generate_variant(
            state,
            temperature=0.5,
            strategy="balanced",
            description="Balanced between literal and natural"
        ))
        
        # Strategy 3: More creative (higher temperature)
        if num_alternatives >= 3:
            alternatives.append(await self._generate_variant(
                state,
                temperature=0.8,
                strategy="creative",
                description="More natural and idiomatic"
            ))
        
        return alternatives
    
    async def _generate_variant(
        self,
        state: Dict,  # Use Dict instead of TranslationState
        temperature: float,
        strategy: str,
        description: str
    ) -> Dict:
        """Generate a single translation variant"""
        
        try:
            # Create LLM with specific temperature
            if hasattr(self.llm, 'model_name'):
                from langchain_openai import ChatOpenAI
                variant_llm = ChatOpenAI(
                    model=self.llm.model_name,
                    temperature=temperature,
                    api_key=self.llm.openai_api_key,
                    timeout=120
                )
            else:
                from langchain_anthropic import ChatAnthropic
                variant_llm = ChatAnthropic(
                    model=self.llm.model,
                    temperature=temperature,
                    anthropic_api_key=self.llm.anthropic_api_key,
                    timeout=120
                )
            
            same_lang = languages_equivalent(state['source_language'], state['target_language'])
            
            if same_lang:
                prompt = f"""Refine this {state['target_language']} text. Strategy: {strategy}.

{LANGUAGE_GUARDRAIL}

Source text:
{state['source_text']}

Provide ONLY the refined text, no explanation."""
            else:
                prompt = f"""Translate from {state['source_language']} to {state['target_language']}. Strategy: {strategy}.

{LANGUAGE_GUARDRAIL}

Source text:
{state['source_text']}

Target audience: {state['target_audience']}
Genre: {state.get('genre', 'General')}

Provide ONLY the translation, no explanation."""
            
            response = await variant_llm.ainvoke([HumanMessage(content=prompt)])
            translation = response.content.strip()
            
            return {
                'strategy': strategy,
                'description': description,
                'temperature': temperature,
                'translation': translation,
                'error': False
            }
        
        except Exception as e:
            return {
                'strategy': strategy,
                'description': description,
                'temperature': temperature,
                'translation': f"Error generating alternative: {str(e)}",
                'error': True
            }
        
# =====================
# SEMANTIC TRANSLATION CACHE
# =====================
class SemanticTranslationCache:
    """Cache translations with semantic similarity-based retrieval"""
    
    def __init__(self, cache_dir: str = ".translation_cache"):
        self.cache_dir = cache_dir
        self.cache = {}  # Full translation cache
        self.sentence_cache = {}  # Sentence-level cache
        self.embeddings = {}  # Store embeddings for similarity search
        self.similarity_threshold = 0.85
        self.embedding_model = None
        self.stats = {
            'hits': 0,
            'misses': 0,
            'partial_hits': 0,
            'similar_hits': 0
        }
        
        # Create cache directory
        os.makedirs(cache_dir, exist_ok=True)
        
        # Load existing cache
        self.load_cache()
    
    def get_embedding_model(self):
        """Lazy load embedding model"""
        if self.embedding_model is None and SENTENCE_TRANSFORMERS_AVAILABLE:
            self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        return self.embedding_model
    
    def get_embedding(self, text: str) -> Optional[np.ndarray]:
        """Get or compute embedding for text"""
        if not SENTENCE_TRANSFORMERS_AVAILABLE:
            return None
        
        text_hash = self.hash_text(text)
        
        if text_hash not in self.embeddings:
            model = self.get_embedding_model()
            if model:
                self.embeddings[text_hash] = model.encode(text)
        
        return self.embeddings.get(text_hash)
    
    def hash_text(self, text: str) -> str:
        """Create hash for text"""
        return hashlib.md5(text.encode()).hexdigest()
    
    def make_cache_key(self, text: str, context: dict) -> str:
        """Create cache key from text and context"""
        context_str = f"{context.get('source_lang')}_{context.get('target_lang')}_{context.get('audience')}_{context.get('genre')}"
        text_hash = self.hash_text(text)
        return f"{text_hash}_{hashlib.md5(context_str.encode()).hexdigest()}"
    
    def context_matches(self, cached_context: dict, query_context: dict) -> bool:
        """Check if contexts are compatible"""
        return (
            cached_context.get('source_lang') == query_context.get('source_lang') and
            cached_context.get('target_lang') == query_context.get('target_lang') and
            cached_context.get('audience') == query_context.get('audience') and
            cached_context.get('genre') == query_context.get('genre')
        )
    
    async def get_cached_translation(
        self, 
        text: str, 
        context: dict
    ) -> Optional[Dict]:
        """Check cache for translation"""
        cache_key = self.make_cache_key(text, context)
        
        # 1. EXACT MATCH (fastest - O(1))
        if cache_key in self.cache:
            self.stats['hits'] += 1
            cached = self.cache[cache_key]
            return {
                'type': 'exact',
                'translation': cached['translation'],
                'metadata': cached['metadata'],
                'speedup': '100x'
            }
        
        # 2. SIMILAR MATCH (still fast - using embeddings)
        if SENTENCE_TRANSFORMERS_AVAILABLE:
            similar = await self.find_similar_translation(text, context)
            if similar:
                self.stats['similar_hits'] += 1
                return {
                    'type': 'similar',
                    'translation': similar['translation'],
                    'similarity': similar['similarity'],
                    'original_text': similar['original_text'],
                    'speedup': '50x',
                    'needs_refinement': True
                }
        
        # 3. SENTENCE-LEVEL CACHE (partial speedup)
        sentence_hits = await self.get_cached_sentences(text, context)
        if sentence_hits and len(sentence_hits['cached']) > len(sentence_hits['all']) * 0.3:
            self.stats['partial_hits'] += 1
            return {
                'type': 'partial',
                'cached_sentences': sentence_hits['cached'],
                'uncached_sentences': sentence_hits['uncached'],
                'speedup': f"{int(len(sentence_hits['cached']) / len(sentence_hits['all']) * 100)}%"
            }
        
        # 4. CACHE MISS
        self.stats['misses'] += 1
        return None
    
    async def find_similar_translation(
        self, 
        text: str, 
        context: dict
    ) -> Optional[Dict]:
        """Find semantically similar cached translation"""
        if not SENTENCE_TRANSFORMERS_AVAILABLE:
            return None
        
        query_embedding = self.get_embedding(text)
        if query_embedding is None:
            return None
        
        best_match = None
        best_similarity = 0.0
        
        for cache_key, cached_data in self.cache.items():
            # Check context compatibility
            if not self.context_matches(cached_data['context'], context):
                continue
            
            # Get cached text embedding
            cached_text = cached_data['source_text']
            cached_embedding = self.get_embedding(cached_text)
            
            if cached_embedding is not None:
                # Compute cosine similarity
                similarity = float(np.dot(query_embedding, cached_embedding) / 
                                 (np.linalg.norm(query_embedding) * np.linalg.norm(cached_embedding)))
                
                if similarity > self.similarity_threshold and similarity > best_similarity:
                    best_similarity = similarity
                    best_match = {
                        'translation': cached_data['translation'],
                        'similarity': similarity,
                        'original_text': cached_text
                    }
        
        return best_match
    
    async def get_cached_sentences(
        self, 
        text: str, 
        context: dict
    ) -> Optional[Dict]:
        """Get cached translations for individual sentences"""
        sentences = [s.strip() + '.' for s in text.split('.') if s.strip()]
        
        cached_sentences = []
        uncached_sentences = []
        
        for i, sentence in enumerate(sentences):
            sent_key = self.make_cache_key(sentence, context)
            
            if sent_key in self.sentence_cache:
                cached_sentences.append({
                    'index': i,
                    'original': sentence,
                    'translation': self.sentence_cache[sent_key]['translation']
                })
            else:
                uncached_sentences.append({
                    'index': i,
                    'original': sentence
                })
        
        if not cached_sentences:
            return None
        
        return {
            'all': sentences,
            'cached': cached_sentences,
            'uncached': uncached_sentences
        }
    
    def store_translation(
        self, 
        text: str, 
        translation: str, 
        context: dict,
        metadata: Optional[dict] = None
    ):
        """Store translation in cache"""
        cache_key = self.make_cache_key(text, context)
        
        # Store full translation
        self.cache[cache_key] = {
            'source_text': text,
            'translation': translation,
            'context': context,
            'metadata': metadata or {},
            'timestamp': datetime.now().isoformat()
        }
        
        # Store sentence-level cache
        sentences = [s.strip() + '.' for s in text.split('.') if s.strip()]
        translated_sentences = [s.strip() + '.' for s in translation.split('.') if s.strip()]
        
        if len(sentences) == len(translated_sentences):
            for src_sent, tgt_sent in zip(sentences, translated_sentences):
                sent_key = self.make_cache_key(src_sent, context)
                self.sentence_cache[sent_key] = {
                    'translation': tgt_sent,
                    'timestamp': datetime.now().isoformat()
                }
        
        # Compute and store embedding
        if SENTENCE_TRANSFORMERS_AVAILABLE:
            self.get_embedding(text)  # This will cache it
        
        # Auto-save periodically
        if len(self.cache) % 10 == 0:
            self.save_cache()
    
    def save_cache(self):
        """Save cache to disk"""
        try:
            # Save main cache
            cache_file = os.path.join(self.cache_dir, 'translations.pkl')
            with open(cache_file, 'wb') as f:
                pickle.dump(self.cache, f)
            
            # Save sentence cache
            sent_cache_file = os.path.join(self.cache_dir, 'sentences.pkl')
            with open(sent_cache_file, 'wb') as f:
                pickle.dump(self.sentence_cache, f)
            
            # Save embeddings
            if self.embeddings:
                emb_file = os.path.join(self.cache_dir, 'embeddings.pkl')
                with open(emb_file, 'wb') as f:
                    pickle.dump(self.embeddings, f)
            
            # Save stats
            stats_file = os.path.join(self.cache_dir, 'stats.json')
            with open(stats_file, 'w') as f:
                json.dump(self.stats, f, indent=2)
            
            return True
        except Exception as e:
            st.warning(f"Cache save failed: {str(e)}")
            return False
    
    def load_cache(self):
        """Load cache from disk"""
        try:
            # Load main cache
            cache_file = os.path.join(self.cache_dir, 'translations.pkl')
            if os.path.exists(cache_file):
                with open(cache_file, 'rb') as f:
                    self.cache = pickle.load(f)
            
            # Load sentence cache
            sent_cache_file = os.path.join(self.cache_dir, 'sentences.pkl')
            if os.path.exists(sent_cache_file):
                with open(sent_cache_file, 'rb') as f:
                    self.sentence_cache = pickle.load(f)
            
            # Load embeddings
            emb_file = os.path.join(self.cache_dir, 'embeddings.pkl')
            if os.path.exists(emb_file):
                with open(emb_file, 'rb') as f:
                    self.embeddings = pickle.load(f)
            
            # Load stats
            stats_file = os.path.join(self.cache_dir, 'stats.json')
            if os.path.exists(stats_file):
                with open(stats_file, 'r') as f:
                    self.stats = json.load(f)
            
            return True
        except Exception as e:
            st.warning(f"Cache load failed: {str(e)}")
            return False
    
    def clear_cache(self):
        """Clear all caches"""
        self.cache = {}
        self.sentence_cache = {}
        self.embeddings = {}
        self.stats = {
            'hits': 0,
            'misses': 0,
            'partial_hits': 0,
            'similar_hits': 0
        }
        
        # Delete cache files
        try:
            for filename in ['translations.pkl', 'sentences.pkl', 'embeddings.pkl', 'stats.json']:
                filepath = os.path.join(self.cache_dir, filename)
                if os.path.exists(filepath):
                    os.remove(filepath)
        except Exception as e:
            st.warning(f"Cache clear failed: {str(e)}")
    
    def get_stats(self) -> dict:
        """Get cache statistics"""
        total_requests = sum(self.stats.values())
        hit_rate = (self.stats['hits'] + self.stats['similar_hits'] + self.stats['partial_hits']) / max(total_requests, 1) * 100
        
        return {
            **self.stats,
            'total_entries': len(self.cache),
            'sentence_entries': len(self.sentence_cache),
            'embedding_entries': len(self.embeddings),
            'hit_rate': hit_rate,
            'total_requests': total_requests
        }
    
    def export_cache(self) -> dict:
        """Export cache for sharing"""
        return {
            'cache': self.cache,
            'sentence_cache': self.sentence_cache,
            'stats': self.stats,
            'exported_at': datetime.now().isoformat()
        }
    
    def import_cache(self, data: dict):
        """Import cache from export"""
        self.cache.update(data.get('cache', {}))
        self.sentence_cache.update(data.get('sentence_cache', {}))
        self.save_cache()

# =====================
# ENTITY TRACKER CLASS
# =====================
import csv
from collections import Counter, defaultdict

class EntitiesTracker:
    """Entity tracking and visualization system with network graphs"""
    
    def __init__(self):
        self.entity_types = {
            'person': {'emoji': '👤', 'color': '#3b82f6'},
            'location': {'emoji': '📍', 'color': '#10b981'},
            'organization': {'emoji': '🏢', 'color': '#f59e0b'},
            'date': {'emoji': '📅', 'color': '#8b5cf6'},
            'custom': {'emoji': '🏷️', 'color': '#ef4444'}
        }
        
        if 'entity_glossary' not in st.session_state:
            st.session_state.entity_glossary = {}
        
        if 'extracted_entities' not in st.session_state:
            st.session_state.extracted_entities = []
    
    def extract_entities(self, text: str) -> List[Dict]:
        """Extract entities from text using glossary and NER patterns"""
        if not text:
            return []
        
        entities = []
        
        # Extract from glossary
        for term, info in st.session_state.entity_glossary.items():
            pattern = re.compile(r'\b' + re.escape(term) + r'\b', re.IGNORECASE)
            matches = pattern.findall(text)
            
            # Check aliases
            alias_count = 0
            for alias in info.get('aliases', []):
                alias_pattern = re.compile(r'\b' + re.escape(alias) + r'\b', re.IGNORECASE)
                alias_matches = alias_pattern.findall(text)
                alias_count += len(alias_matches)
            
            total_count = len(matches) + alias_count
            
            if total_count > 0:
                entities.append({
                    'name': term,
                    'type': info.get('type', 'custom'),
                    'count': total_count,
                    'description': info.get('description', ''),
                    'from_glossary': True
                })
        
        # Auto-detection patterns
        # Person names
        person_pattern = r'\b([A-Z][a-z]+ [A-Z][a-z]+)\b'
        for match in re.finditer(person_pattern, text):
            name = match.group(1)
            if name not in [e['name'] for e in entities]:
                count = len(re.findall(r'\b' + re.escape(name) + r'\b', text))
                entities.append({
                    'name': name,
                    'type': 'person',
                    'count': count,
                    'description': 'Auto-detected',
                    'auto_detected': True
                })
        
        # Organizations
        org_pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Inc|Corp|LLC|Ltd|Company|Group|Foundation))\b'
        for match in re.finditer(org_pattern, text):
            org = match.group(1)
            if org not in [e['name'] for e in entities]:
                count = len(re.findall(r'\b' + re.escape(org) + r'\b', text))
                entities.append({
                    'name': org,
                    'type': 'organization',
                    'count': count,
                    'description': 'Auto-detected',
                    'auto_detected': True
                })
        
        return entities
    
    def upload_glossary(self, file) -> bool:
        """Process uploaded glossary file"""
        try:
            content = file.read()
            
            if file.name.endswith('.json'):
                new_glossary = json.loads(content)
            elif file.name.endswith('.csv'):
                csv_data = io.StringIO(content.decode('utf-8'))
                reader = csv.DictReader(csv_data)
                new_glossary = {}
                for row in reader:
                    term = row.get('term', '').strip()
                    if term:
                        new_glossary[term] = {
                            'type': row.get('type', 'custom'),
                            'description': row.get('description', ''),
                            'aliases': [a.strip() for a in row.get('aliases', '').split(',') if a.strip()]
                        }
            else:
                lines = content.decode('utf-8').split('\n')
                new_glossary = {}
                for line in lines:
                    term = line.strip()
                    if term:
                        new_glossary[term] = {'type': 'custom', 'description': '', 'aliases': []}
            
            st.session_state.entity_glossary.update(new_glossary)
            return True
        except Exception as e:
            st.error(f"Error processing glossary: {str(e)}")
            return False
    
    def visualize_network(self, entities: List[Dict]):
        """Create network visualization using plotly"""
        if not NETWORKX_AVAILABLE or not PLOTLY_AVAILABLE:
            st.warning("Install networkx and plotly for network visualization")
            return
        
        # Create graph
        G = nx.Graph()
        
        # Add nodes
        for entity in entities:
            G.add_node(entity['name'],
                      type=entity['type'],
                      count=entity['count'],
                      description=entity.get('description', ''))
        
        # Add edges (simplified co-occurrence)
        for i, e1 in enumerate(entities):
            for e2 in entities[i+1:]:
                weight = min(e1['count'], e2['count'])
                G.add_edge(e1['name'], e2['name'], weight=weight)
        
        # Create layout
        pos = nx.spring_layout(G, k=2, iterations=50)
        
        # Create edge traces
        edge_traces = []
        for edge in G.edges(data=True):
            x0, y0 = pos[edge[0]]
            x1, y1 = pos[edge[1]]
            trace = go.Scatter(x=[x0, x1, None], y=[y0, y1, None],
                             mode='lines',
                             line=dict(width=0.5, color='#888'),
                             hoverinfo='none')
            edge_traces.append(trace)
        
        # Create node trace
        node_x = []
        node_y = []
        node_text = []
        node_color = []
        node_size = []
        
        for node in G.nodes():
            x, y = pos[node]
            node_x.append(x)
            node_y.append(y)
            
            entity = next(e for e in entities if e['name'] == node)
            node_text.append(f"{node}<br>Type: {entity['type']}<br>Count: {entity['count']}")
            node_color.append(self.entity_types[entity['type']]['color'])
            node_size.append(10 + min(entity['count'] * 3, 50))
        
        node_trace = go.Scatter(
            x=node_x, y=node_y,
            mode='markers+text',
            text=[n for n in G.nodes()],
            textposition="top center",
            hoverinfo='text',
            hovertext=node_text,
            marker=dict(
                size=node_size,
                color=node_color,
                line=dict(width=2, color='white')
            )
        )
        
        # Create figure
        fig = go.Figure(data=edge_traces + [node_trace],
                       layout=go.Layout(
                           title='Entity Network',
                           showlegend=False,
                           hovermode='closest',
                           margin=dict(b=0,l=0,r=0,t=40),
                           xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                           yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                           height=600
                       ))
        
        st.plotly_chart(fig, use_container_width=True)

# ---- Helpers for same-language paths ----
def normalize_lang_label(label: str) -> str:
    """Normalize a language label like 'English (US)' -> 'english'."""
    core = label.split("(")[0].strip().lower()
    return core

def languages_equivalent(src: str, tgt: str) -> bool:
    return normalize_lang_label(src) == normalize_lang_label(tgt)

def mode_phrase(src: str, tgt: str) -> str:
    """Human-readable short descriptor for UI headers."""
    if languages_equivalent(src, tgt):
        return f"{tgt} – refine"
    return f"{src} → {tgt}"

def split_notes(content: str, labels: List[str]) -> (str, Optional[str], Optional[str]):
    """
    Split content on the first matching label in labels list (e.g., ["TRANSLATOR NOTES:", "EDITOR NOTES:"]).
    Returns (main_text, found_label, notes_text).
    """
    for lab in labels:
        if lab in content:
            parts = content.split(lab, 1)
            return parts[0].strip(), lab, parts[1].strip()
    return content.strip(), None, None

# LangSmith integration (optional)
try:
    from langsmith import Client
    LANGSMITH_AVAILABLE = True
except ImportError:
    LANGSMITH_AVAILABLE = False

# Model imports
try:
    from langchain_openai import ChatOpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

try:
    from langchain_anthropic import ChatAnthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

# Document processing
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# BERTScore (optional; only used in same-language mode)
try:
    from bert_score import score as bert_score_fn
    BERT_AVAILABLE = True
except Exception:
    BERT_AVAILABLE = False

# Checkpointer
from langgraph.checkpoint.memory import MemorySaver

# Caching
from langchain_core.caches import InMemoryCache
from langchain_core.globals import set_llm_cache
set_llm_cache(InMemoryCache())

# =====================
# AGENT ROLE DESCRIPTIONS
# =====================
AGENT_DESCRIPTIONS = {
    "planning": {
        "name": "Workflow Planning Specialist",
        "emoji": "📋",
        "role": "Intelligent Agent Selection & Routing",
        "description": """
        **Primary Responsibility:** Analyze source text and dynamically determine which agents are needed.
        **Key Functions:**
        - Assess text complexity and characteristics
        - Detect technical, cultural, and literary elements
        - Select optimal agent sequence for the task
        - Estimate time and cost savings
        - Provide reasoning for all decisions
        """,
        "expertise": ["Text analysis", "Workflow optimization", "Resource efficiency", "Quality assurance"]
    },
    "literal_translator": {
        "name": "Baseline Specialist",
        "emoji": "🔤",
        "role": "Foundational Accuracy (or faithful baseline in refine mode)",
        "description": """
        **Primary Responsibility:** Provides the baseline pass:
        - Translation mode: create an accurate word-for-word draft.
        - Refine mode (same-language): perform a faithful copy-edit that preserves meaning and structure.
        **Key Functions:**
        - Preserve semantic content and important structure
        - Identify idioms, ambiguities, cultural expressions
        - Preserve technical terms and proper nouns
        - Document challenges for downstream agents
        """,
        "expertise": ["Lexical mapping", "Semantic preservation", "Grammatical structure", "Technical terminology"]
    },
    "cultural_adapter": {
        "name": "Cultural Localization Expert",
        "emoji": "🌍",
        "role": "Cross-Cultural Bridge (or register-localization in refine mode)",
        "description": """
        **Primary Responsibility:** Make content culturally natural in target context.
        - In refine mode, ensure register and references suit the target variety (e.g., US vs UK English).
        """,
        "expertise": ["Localization", "Idiomatic expressions", "Cross-cultural communication", "Audience adaptation"]
    },
    "tone_specialist": {
        "name": "Tone & Voice Consistency Director",
        "emoji": "🎭",
        "role": "Stylistic Harmony & Readability",
        "description": "Ensure consistent tone, voice, pacing, and readability.",
        "expertise": ["Stylistics", "Rhetorical analysis", "Reader psychology", "Narrative voice"]
    },
    "technical_reviewer": {
        "name": "Technical Accuracy Auditor",
        "emoji": "🔬",
        "role": "Precision & Factual Integrity",
        "description": "Verify notation, terminology, measurements, and formats.",
        "expertise": ["Technical writing", "Scientific notation", "Terminology", "Quality assurance"]
    },
    "literary_editor": {
        "name": "Literary Style & Excellence Editor",
        "emoji": "✍️",
        "role": "Publication-Ready Prose Crafting",
        "description": "Elevate prose to publication quality without altering meaning.",
        "expertise": ["Literary criticism", "Creative writing", "Stylistics", "Publishing standards"]
    },
    "quality_controller": {
        "name": "Master Quality Synthesizer",
        "emoji": "✅",
        "role": "Final Integration & Excellence Assurance",
        "description": "Integrate all contributions and approve the final version.",
        "expertise": ["Editorial judgment", "Synthesis", "Holistic QA"]
    },
    "bertscore_validator": {
        "name": "Semantic Fidelity Validator",
        "emoji": "🎯",
        "role": "Semantic Similarity Assurance (same-language only)",
        "description": """
        **Primary Responsibility:** Ensure semantic fidelity in same-language refinement.
        - Only active when source and target languages are equivalent
        - Validates BERTScore ≥ 0.8 for semantic preservation
        - Iteratively refines output to meet threshold
        **Key Functions:**
        - Compute BERTScore (P/R/F1) against source
        - Identify semantic drift areas
        - Apply targeted refinements without over-editing
        - Preserve meaning while improving clarity
        """,
        "expertise": ["Semantic analysis", "Embedding comparison", "Iterative refinement", "Fidelity validation"]
    }
}

# =====================
# State Definition (UPDATED WITH NEW FIELDS)
# =====================
class TranslationState(TypedDict):
    source_text: str
    source_language: str
    target_language: str
    target_audience: str
    genre: str
    # Versions
    literal_translation: str
    cultural_adaptation: str
    tone_adjustment: str
    technical_review_version: str
    literary_polish: str
    final_translation: str
    # Issues
    literal_issues: List[Dict]
    cultural_issues: List[Dict]
    tone_issues: List[Dict]
    technical_issues: List[Dict]
    literary_issues: List[Dict]
    # Critical passages
    critical_passages: List[Dict]
    # Workflow tracking
    agent_notes: Annotated[List[str], operator.add]
    agent_decisions: List[Dict]
    human_feedback: Optional[str]
    revision_count: int
    needs_human_review: bool
    # Entity tracking (Optional)
    source_entities: Optional[List[Dict]]
    translated_entities: Optional[List[Dict]]
    entity_preservation_rate: Optional[float]
    # BERTScore tracking
    bertscore_attempts: int
    bertscore_history: List[Dict]
    # Planning fields
    agent_plan: List[str]
    agent_plan_reasoning: Dict[str, str]
    planning_analysis: Dict[str, any]
    skipped_agents: List[str]
    estimated_complexity: str
    current_agent_index: int
    planning_enabled: bool
    # Cache tracking
    cache_hit: Optional[str]
    cache_speedup: Optional[str]
    # NEW: Confidence scores
    confidence_scores: Optional[Dict[str, float]]
    # Metadata
    started_at: str
    completed_at: Optional[str]

# =====================
# File Processing Utilities
# =====================
def read_uploaded_file(uploaded_file) -> str:
    """Read content from uploaded file"""
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        if file_extension in ('txt', 'md'):
            return uploaded_file.read().decode('utf-8')
        elif file_extension == 'docx':
            if not DOCX_AVAILABLE:
                st.error("python-docx not installed. Install with: pip install python-docx")
                return ""
            doc = Document(uploaded_file)
            return '\n\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            st.error(f"Unsupported file type: {file_extension}")
            return ""
    except Exception as e:
        st.error(f"Error reading file: {str(e)}")
        return ""

def create_docx_file(text: str, title: str = "Translation") -> io.BytesIO:
    """Create a formatted Word document"""
    if not DOCX_AVAILABLE:
        st.error("python-docx not installed. Install with: pip install python-docx")
        return None
    
    doc = Document()
    title_paragraph = doc.add_heading(title, level=1)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    timestamp = doc.add_paragraph()
    timestamp.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    for para in text.split('\n\n'):
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def create_markdown_file(text: str, title: str = "Translation") -> str:
    return f"""# {title}

**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

---

{text}
"""

def extract_critical_passages(text: str, issues_list: List[Dict]) -> List[Dict]:
    """Extract critical passages that need review based on agent feedback"""
    critical = []
    for issue in issues_list:
        content = issue.get('content', '')
        issue_type = issue.get('type', '')
        agent = issue.get('agent', '')
        
        if any(k in content.lower() for k in ['critical', 'warning', 'error', 'ambiguous', 'unclear', 'needs review']):
            sentences = text.split('.')
            for i, sentence in enumerate(sentences):
                if len(sentence.strip()) > 10:
                    critical.append({
                        'passage': sentence.strip() + '.',
                        'issue': content[:200],
                        'agent': agent,
                        'type': issue_type,
                        'sentence_index': i
                    })
    
    return critical[:10]

# =====================
# Optional: Language Reinforcement
# =====================
async def reinforce_language(llm: BaseChatModel, text: str, target_language: str) -> str:
    """Lightweight fixer to ensure output remains in the target language only."""
    try:
        resp = await llm.ainvoke([
            SystemMessage(content="Ensure the following text is in the specified language only, without code-switching."),
            HumanMessage(content=f"{LANGUAGE_GUARDRAIL}\n\nTarget language: {target_language}\n\nText:\n{text}\n\nReturn the corrected text in {target_language}, with no extra commentary.")
        ])
        return resp.content.strip()
    except Exception:
        return text

# =====================
# BERTScore Utility (same-language only)
# =====================
def compute_bertscore(candidate: str, reference: str) -> Optional[Dict[str, float]]:
    """Compute BERTScore (P/R/F1) if package is available; returns None otherwise."""
    if not BERT_AVAILABLE:
        return None
    try:
        P, R, F1 = bert_score_fn([candidate], [reference], lang="en", rescale_with_baseline=True)
        return {
            "precision": float(P.mean().item()),
            "recall": float(R.mean().item()),
            "f1": float(F1.mean().item()),
        }
    except Exception:
        return None

# =====================
# Visualization helpers
# =====================
def sentence_lengths(text: str) -> List[int]:
    """Return per-sentence word counts for a text."""
    if not text:
        return []
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    return [len(s.split()) for s in sentences if s.strip()]

def nonzero_bins(max_len: int) -> List[int]:
    """Nice 5-word bins for histograms."""
    step = 5
    upper = max(5, ((max_len + step - 1) // step) * step)
    return list(range(0, upper + step, step))

_WORD_RE = re.compile(r"[A-Za-zÀ-ÖØ-öø-ÿ']+")

def tokenize_words(text: str) -> List[str]:
    if not text:
        return []
    return [w.lower() for w in _WORD_RE.findall(text)]

def word_frequencies(text: str, stopwords: Optional[set] = None) -> Dict[str, int]:
    tokens = tokenize_words(text)
    if stopwords is None:
        stopwords = set()
    cnt = collections.Counter(w for w in tokens if w not in stopwords and len(w) > 1)
    return dict(cnt)

def render_wordcloud_from_freq(freqs: Dict[str, int], title: str):
    if not freqs:
        st.info(f"No tokens to display for **{title}**.")
        return
    
    wc = WordCloud(
        width=1000,
        height=500,
        background_color="white",
        collocations=False
    ).generate_from_frequencies(freqs)
    
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.imshow(wc, interpolation="bilinear")
    ax.axis("off")
    ax.set_title(title)
    st.pyplot(fig)

def text_similarity(a: str, b: str) -> float:
    return difflib.SequenceMatcher(None, a or "", b or "").ratio()

# =====================
# PLANNING AGENT
# =====================
class PlanningAgent:
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
        self.name = "Workflow Planning Specialist"
        self.emoji = "📋"
    
    async def analyze_and_plan(self, state: TranslationState) -> TranslationState:
        """Analyze text and create execution plan"""
        if state.get('agent_plan'):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Plan already exists - skipping")
            return state
        
        attempt = 0
        max_attempts = 2
        
        while attempt < max_attempts:
            try:
                # Check if planning is disabled (manual override)
                if not state.get('planning_enabled', True):
                    state['agent_plan'] = [
                        'literal_translator', 'cultural_adapter', 'tone_specialist',
                        'technical_reviewer', 'literary_editor', 'finalize', 'bertscore_validator'
                    ]
                    state['agent_plan_reasoning'] = {
                        'literal_translator': '✅ Required - baseline',
                        'cultural_adapter': '✅ Included - planning disabled',
                        'tone_specialist': '✅ Included - planning disabled',
                        'technical_reviewer': '✅ Included - planning disabled',
                        'literary_editor': '✅ Included - planning disabled',
                        'finalize': '✅ Required - final review',
                        'bertscore_validator': '✅ Included - planning disabled'
                    }
                    state['skipped_agents'] = []
                    state['estimated_complexity'] = 'full'
                    state['current_agent_index'] = 0
                    state['planning_analysis'] = {'mode': 'manual_override'}
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Planning disabled - running all agents")
                    break
                
                source_text = state['source_text']
                source_lang = state['source_language']
                target_lang = state['target_language']
                audience = state['target_audience']
                genre = state.get('genre', 'General')
                same_lang = languages_equivalent(source_lang, target_lang)
                
                # Check for manual overrides from session state
                force_cultural = st.session_state.get('force_cultural_adapter', False)
                force_tone = st.session_state.get('force_tone_specialist', False)
                force_technical = st.session_state.get('force_technical_reviewer', False)
                force_literary = st.session_state.get('force_literary_editor', False)
                force_bertscore = st.session_state.get('force_bertscore_validator', False)
                
                system_prompt = f"""You are an expert translation workflow planner. Analyze the source text and determine which translation agents are needed.

AVAILABLE AGENTS:
1. literal_translator (ALWAYS REQUIRED) - Creates initial translation/refinement
2. cultural_adapter - Handles idioms, cultural references, localization
3. tone_specialist - Ensures voice consistency and readability
4. technical_reviewer - Validates technical accuracy, terminology, units
5. literary_editor - Elevates prose quality for publication
6. finalize (ALWAYS REQUIRED) - Final quality synthesis
7. bertscore_validator - Semantic fidelity validation (same-language only)

ANALYSIS CRITERIA:
- Text length: {len(source_text)} characters, ~{len(source_text.split())} words
- Language pair: {source_lang} → {target_lang} (same-language: {same_lang})
- Target audience: {audience}
- Genre: {genre}
- Look for: technical terms, cultural references, idioms, literary elements, tone consistency issues

SELECTION GUIDELINES:
- literal_translator & finalize: ALWAYS include
- cultural_adapter: Include if cross-language OR idioms/cultural content OR same-lang with register differences
- tone_specialist: Include if >500 words OR inconsistent tone detected OR multiple sections
- technical_reviewer: Include if technical terminology, formulas, units, citations present
- literary_editor: Include if literary genre OR narrative elements OR high-polish audience
- bertscore_validator: ALWAYS include if same-language mode. DO NOT SKIP.

EFFICIENCY GOAL: Include only agents that will meaningfully improve output. When uncertain, include the agent (favor quality over speed).

OUTPUT FORMAT: Return ONLY valid JSON (no markdown, no backticks):
{{
    "required_agents": ["literal_translator", "cultural_adapter", "finalize"],
    "reasoning": {{
        "literal_translator": "✅ Required - baseline",
        "cultural_adapter": "✅ Detected 2 idioms and cross-language pair",
        "tone_specialist": "⏭️ Skipped - short text (120 words), consistent tone",
        "technical_reviewer": "⏭️ Skipped - no technical content",
        "literary_editor": "⏭️ Skipped - casual genre",
        "finalize": "✅ Required - final review",
        "bertscore_validator": "⏭️ Skipped - cross-language"
    }},
    "analysis": {{
        "length": 120,
        "complexity_score": 2.3,
        "technical_terms_count": 0,
        "cultural_references_count": 2,
        "tone_consistency": "high",
        "literary_elements": "minimal"
    }},
    "estimated_complexity": "simple"
}}
"""
                
                user_prompt = f"""Analyze this text and create an optimal translation workflow plan:

SOURCE TEXT:
{source_text[:2000]}{'...' if len(source_text) > 2000 else ''}

Return ONLY the JSON plan, no other text."""
                
                response = await self.llm.ainvoke([
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=user_prompt)
                ])
                
                content = response.content.strip()
                content = re.sub(r'```json\s*', '', content)
                content = re.sub(r'```\s*$', '', content)
                
                plan_data = json.loads(content)
                
                required_agents = plan_data.get('required_agents', [])
                reasoning = plan_data.get('reasoning', {})
                analysis = plan_data.get('analysis', {})
                complexity = plan_data.get('estimated_complexity', 'moderate')
                
                # Apply manual overrides
                if force_cultural and 'cultural_adapter' not in required_agents:
                    required_agents.insert(1, 'cultural_adapter')
                    reasoning['cultural_adapter'] = '✅ Forced by user override'
                
                if force_tone and 'tone_specialist' not in required_agents:
                    insert_pos = 2 if 'cultural_adapter' in required_agents else 1
                    required_agents.insert(insert_pos, 'tone_specialist')
                    reasoning['tone_specialist'] = '✅ Forced by user override'
                
                if force_technical and 'technical_reviewer' not in required_agents:
                    finalize_pos = required_agents.index('finalize') if 'finalize' in required_agents else len(required_agents)
                    required_agents.insert(finalize_pos, 'technical_reviewer')
                    reasoning['technical_reviewer'] = '✅ Forced by user override'
                
                if force_literary and 'literary_editor' not in required_agents:
                    finalize_pos = required_agents.index('finalize') if 'finalize' in required_agents else len(required_agents)
                    required_agents.insert(finalize_pos, 'literary_editor')
                    reasoning['literary_editor'] = '✅ Forced by user override'
                
                if force_bertscore and 'bertscore_validator' not in required_agents and same_lang:
                    required_agents.append('bertscore_validator')
                    reasoning['bertscore_validator'] = '✅ Forced by user override'
                
                # Ensure required agents are present
                if 'literal_translator' not in required_agents:
                    required_agents.insert(0, 'literal_translator')
                if 'finalize' not in required_agents:
                    required_agents.append('finalize')
                
                # Enforce bertscore_validator always runs in same-language mode
                if same_lang and BERT_AVAILABLE and 'bertscore_validator' not in required_agents:
                    required_agents.append('bertscore_validator')
                    reasoning['bertscore_validator'] = '✅ Always required in same-language mode'
                
                all_agents = ['literal_translator', 'cultural_adapter', 'tone_specialist',
                             'technical_reviewer', 'literary_editor', 'finalize', 'bertscore_validator']
                skipped = [a for a in all_agents if a not in required_agents]
                
                state['agent_plan'] = required_agents
                state['agent_plan_reasoning'] = reasoning
                state['planning_analysis'] = analysis
                state['skipped_agents'] = skipped
                state['estimated_complexity'] = complexity
                state['current_agent_index'] = 0
                
                total_agents = len(all_agents)
                used_agents = len(required_agents)
                savings_pct = int((1 - used_agents / total_agents) * 100)
                
                state['agent_notes'].append(
                    f"{self.emoji} {self.name}: Plan created - {used_agents}/{total_agents} agents "
                    f"({savings_pct}% savings) - Complexity: {complexity}"
                )
                break
            
            except Exception as e:
                attempt += 1
                if attempt < max_attempts:
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Retry attempt {attempt} after error: {str(e)}")
                else:
                    st.warning(f"Planning agent failed after {max_attempts} attempts, using conservative defaults: {str(e)}")
                    state['agent_plan'] = [
                        'literal_translator', 'cultural_adapter', 'tone_specialist',
                        'technical_reviewer', 'literary_editor', 'finalize'
                    ]
                    if same_lang and BERT_AVAILABLE:
                        state['agent_plan'].append('bertscore_validator')
                    
                    state['agent_plan_reasoning'] = {
                        a: '✅ Included (fallback mode)' for a in state['agent_plan']
                    }
                    state['skipped_agents'] = []
                    state['estimated_complexity'] = 'unknown'
                    state['current_agent_index'] = 0
                    state['planning_analysis'] = {'error': str(e), 'mode': 'fallback'}
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Planning failed - using all agents (safe mode)")
        
        return state

# =====================
# FIXED ROUTER FUNCTION
# =====================
def route_to_next_agent(state: TranslationState) -> str:
    """Determine next agent based on the plan"""
    agent_plan = state.get('agent_plan', [])
    current_index = state.get('current_agent_index', 0)
    
    if current_index >= len(agent_plan):
        return END
    
    next_agent = agent_plan[current_index]
    
    return next_agent

# =====================
# AGENT CLASSES - Each increments index
# =====================
class LiteralTranslationAgent:
    def __init__(self, llm: BaseChatModel, cache: SemanticTranslationCache):
        self.llm = llm
        self.cache = cache
        self.name = "Baseline Specialist"
        self.emoji = "🔤"
    
    async def translate(self, state: TranslationState) -> TranslationState:
        if state.get('literal_translation'):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Already complete - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        # Check cache first
        context = {
            'source_lang': state['source_language'],
            'target_lang': state['target_language'],
            'audience': state['target_audience'],
            'genre': state.get('genre', 'General')
        }
        
        cached = await self.cache.get_cached_translation(state['source_text'], context)
        
        if cached and cached['type'] == 'exact':
            state['literal_translation'] = cached['translation']
            state['literal_issues'] = []
            state['cache_hit'] = 'exact'
            state['cache_speedup'] = cached['speedup']
            state['agent_notes'].append(f"⚡ {self.emoji} {self.name}: CACHE HIT (exact match, {cached['speedup']} faster)")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        if cached and cached['type'] == 'similar':
            state['agent_notes'].append(f"⚡ {self.emoji} {self.name}: CACHE HIT (similar match, {cached['similarity']:.2f} similarity)")
            # Use similar translation as starting point
            base_translation = cached['translation']
            state['cache_hit'] = 'similar'
            state['cache_speedup'] = cached['speedup']
        else:
            base_translation = None
        
        attempt = 0
        max_attempts = 2
        
        while attempt < max_attempts:
            try:
                same_lang = languages_equivalent(state['source_language'], state['target_language'])
                system_prompt = (
                    "You provide the baseline pass:\n"
                    "- Translation mode: accurate, faithful literal translation.\n"
                    "- Refine mode: faithful copy-edit that preserves meaning and structure.\n\n"
                    + LANGUAGE_GUARDRAIL
                )
                
                if same_lang:
                    user_prompt = f"""Refine the following text in {state['target_language']} without changing its meaning.

{LANGUAGE_GUARDRAIL}

**GOAL (REFINE MODE):**
- Preserve semantics, improve clarity and microstructure
- Flag idioms/ambiguities for downstream agents
- Keep terminology and proper nouns intact
- OUTPUT FORMAT: Provide the refined text first, then add "EDITOR NOTES:" (notes must be in {state['target_language']}).

**TEXT TO REFINE:**
{state['source_text']}

**AUDIENCE**: {state['target_audience']}
**GENRE**: {state.get('genre', 'General')}
"""
                    
                    if 'enable_entity_awareness' in st.session_state and st.session_state.enable_entity_awareness:
                        if state.get('source_entities'):
                            entity_list = "\n".join([f"- {e['name']} ({e['type']})" for e in state['source_entities'][:15]])
                            user_prompt += f"\n\n**IMPORTANT ENTITIES TO PRESERVE:**\n{entity_list}\n"
                else:
                    user_prompt = f"""Translate the following text from {state['source_language']} to {state['target_language']}.

{LANGUAGE_GUARDRAIL}

**CRITICAL INSTRUCTIONS:**
1. Translate with maximum fidelity to the original meaning.
2. Maintain sentence structure initially.
3. Flag idioms/ambiguities for downstream agents.
4. OUTPUT FORMAT: Provide the literal translation first, then "TRANSLATOR NOTES:" (notes must be in {state['target_language']}).

**SOURCE TEXT:**
{state['source_text']}

**TARGET AUDIENCE**: {state['target_audience']}
**GENRE**: {state.get('genre', 'General')}
"""
                    
                    if 'enable_entity_awareness' in st.session_state and st.session_state.enable_entity_awareness:
                        if state.get('source_entities'):
                            entity_list = "\n".join([f"- {e['name']} ({e['type']})" for e in state['source_entities'][:15]])
                            user_prompt += f"\n\n**IMPORTANT ENTITIES TO PRESERVE:**\n{entity_list}\n"
                
                response = await self.llm.ainvoke([SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)])
                content = response.content
                
                translation, found_label, notes = split_notes(content, ["TRANSLATOR NOTES:", "EDITOR NOTES:"])
                
                issues = []
                if notes:
                    issues.append({
                        "agent": self.name,
                        "type": (found_label[:-1].lower().replace(" ", "_") if found_label else "notes"),
                        "content": notes
                    })
                
                translation = await reinforce_language(self.llm, translation, state['target_language'])
                
                state['literal_translation'] = translation
                state['literal_issues'] = issues
                
                # Store in cache
                self.cache.store_translation(
                    state['source_text'],
                    translation,
                    context,
                    {'agent': self.name, 'timestamp': datetime.now().isoformat()}
                )
                
                state['agent_notes'].append(f"{self.emoji} {self.name}: Baseline {'refinement' if same_lang else 'literal translation'} complete")
                break
            
            except Exception as e:
                attempt += 1
                if attempt < max_attempts:
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Retry attempt {attempt} after error: {str(e)}")
                else:
                    st.error(f"Error in Literal step after {max_attempts} attempts: {str(e)}")
                    state['literal_translation'] = state['source_text']
                    state['literal_issues'] = [{"agent": self.name, "type": "error", "content": str(e)}]
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Error - using source text")
        
        # INCREMENT INDEX
        state['current_agent_index'] = state.get('current_agent_index', 0) + 1
        return state

class CulturalAdaptationAgent:
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
        self.name = "Cultural Localization Expert"
        self.emoji = "🌍"
    
    async def adapt(self, state: TranslationState) -> TranslationState:
        if state.get('cultural_adaptation'):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Already complete - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        previous_text = state.get('literal_translation') or ""
        
        if text_similarity(previous_text, state.get('cultural_adaptation', '')) > 0.95:
            state['agent_notes'].append(f"{self.emoji} {self.name}: Minimal changes detected - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        attempt = 0
        max_attempts = 2
        
        while attempt < max_attempts:
            try:
                same_lang = languages_equivalent(state['source_language'], state['target_language'])
                system_prompt = "You adapt content for cultural/register naturalness in the target context.\n\n" + LANGUAGE_GUARDRAIL
                
                if same_lang:
                    task_text = (
                        "Refine for the target variety's norms (e.g., spelling, idioms, punctuation, register). "
                        "Replace region-specific expressions with appropriate target-variety equivalents."
                    )
                    notes_label = "EDITOR NOTES:"
                else:
                    task_text = (
                        "Replace source-culture idioms with target equivalents, adapt references, "
                        "and adjust communication style for the target audience."
                    )
                    notes_label = "CULTURAL NOTES:"
                
                user_prompt = f"""Work on the following text.

{LANGUAGE_GUARDRAIL}

**TARGET CONTEXT**: {state['target_language']} / Audience: {state['target_audience']}

**TEXT:**
{state['literal_translation']}

**YOUR TASKS:**
- {task_text}

**OUTPUT**: Provide the adapted text, then {notes_label} (notes must be in {state['target_language']})."""
                
                response = await self.llm.ainvoke([SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)])
                content = response.content
                
                adapted, found_label, notes = split_notes(content, ["CULTURAL NOTES:", "EDITOR NOTES:"])
                
                issues = []
                if notes:
                    issues.append({
                        "agent": self.name,
                        "type": (found_label[:-1].lower().replace(" ", "_") if found_label else "notes"),
                        "content": notes
                    })
                
                adapted = await reinforce_language(self.llm, adapted, state['target_language'])
                
                state['cultural_adaptation'] = adapted
                state['cultural_issues'] = issues
                state['agent_notes'].append(f"{self.emoji} {self.name}: {'Register' if same_lang else 'Cultural'} adaptation complete")
                break
            
            except Exception as e:
                attempt += 1
                if attempt < max_attempts:
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Retry attempt {attempt} after error: {str(e)}")
                else:
                    st.error(f"Error in Cultural step after {max_attempts} attempts: {str(e)}")
                    state['cultural_adaptation'] = state['literal_translation']
                    state['cultural_issues'] = [{"agent": self.name, "type": "error", "content": str(e)}]
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Error - using previous version")
        
        # INCREMENT INDEX
        state['current_agent_index'] = state.get('current_agent_index', 0) + 1
        return state

class ToneConsistencyAgent:
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
        self.name = "Tone & Voice Consistency Director"
        self.emoji = "🎭"
    
    async def adjust_tone(self, state: TranslationState) -> TranslationState:
        if state.get('tone_adjustment'):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Already complete - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        previous_text = state.get('cultural_adaptation') or state.get('literal_translation') or ""
        
        if text_similarity(previous_text, state.get('tone_adjustment', '')) > 0.95:
            state['agent_notes'].append(f"{self.emoji} {self.name}: Minimal changes detected - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        attempt = 0
        max_attempts = 2
        
        while attempt < max_attempts:
            try:
                system_prompt = "You ensure tone/voice consistency and natural readability.\n\n" + LANGUAGE_GUARDRAIL
                
                user_prompt = f"""Adjust this text for tone consistency and readability.

{LANGUAGE_GUARDRAIL}

**AUDIENCE**: {state['target_audience']}

**TEXT:**
{previous_text}

**YOUR TASKS:**
1. Vary sentence length for rhythm
2. Match formality to audience
3. Ensure consistent voice
4. Optimize readability

**OUTPUT**: Provide the adjusted text, then "TONE NOTES:" (notes must be in {state['target_language']})."""
                
                response = await self.llm.ainvoke([SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)])
                content = response.content
                
                adjusted, found_label, notes = split_notes(content, ["TONE NOTES:"])
                
                issues = []
                if notes:
                    issues.append({"agent": self.name, "type": "tone_notes", "content": notes})
                
                adjusted = await reinforce_language(self.llm, adjusted, state['target_language'])
                
                state['tone_adjustment'] = adjusted
                state['tone_issues'] = issues
                state['agent_notes'].append(f"{self.emoji} {self.name}: Tone/readability optimized")
                break
            
            except Exception as e:
                attempt += 1
                if attempt < max_attempts:
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Retry attempt {attempt} after error: {str(e)}")
                else:
                    st.error(f"Error in Tone step after {max_attempts} attempts: {str(e)}")
                    state['tone_adjustment'] = previous_text
                    state['tone_issues'] = [{"agent": self.name, "type": "error", "content": str(e)}]
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Error - using previous version")
        
        # INCREMENT INDEX
        state['current_agent_index'] = state.get('current_agent_index', 0) + 1
        return state

class TechnicalReviewAgent:
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
        self.name = "Technical Accuracy Auditor"
        self.emoji = "🔬"
    
    async def review(self, state: TranslationState) -> TranslationState:
        if state.get('technical_review_version'):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Already complete - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        previous_text = (state.get('tone_adjustment') or
                         state.get('cultural_adaptation') or
                         state.get('literal_translation') or "")
        
        if text_similarity(previous_text, state.get('technical_review_version', '')) > 0.95:
            state['agent_notes'].append(f"{self.emoji} {self.name}: Minimal changes detected - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        attempt = 0
        max_attempts = 2
        
        while attempt < max_attempts:
            try:
                system_prompt = "You verify notation, terminology, units, and formatting.\n\n" + LANGUAGE_GUARDRAIL
                
                user_prompt = f"""Review the following for technical accuracy.

{LANGUAGE_GUARDRAIL}

**TEXT:**
{previous_text}

**YOUR TASKS:**
1. Verify notation/symbols
2. Check terminology
3. Validate measurements/units
4. Ensure number/date/time formatting

**OUTPUT**: Provide the reviewed text, then "TECHNICAL NOTES:" if corrections made (notes must be in {state['target_language']})."""
                
                response = await self.llm.ainvoke([SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)])
                content = response.content
                
                reviewed, found_label, notes = split_notes(content, ["TECHNICAL NOTES:"])
                
                issues = []
                if notes:
                    issues.append({"agent": self.name, "type": "technical_notes", "content": notes})
                
                needs_review = "NEEDS_REVIEW" in content or "NEEDS REVIEW" in content
                
                reviewed = await reinforce_language(self.llm, reviewed, state['target_language'])
                
                state['technical_review_version'] = reviewed
                state['technical_issues'] = issues
                state['needs_human_review'] = needs_review
                state['agent_notes'].append(f"{self.emoji} {self.name}: Technical review completed")
                break
            
            except Exception as e:
                attempt += 1
                if attempt < max_attempts:
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Retry attempt {attempt} after error: {str(e)}")
                else:
                    st.error(f"Error in Technical step after {max_attempts} attempts: {str(e)}")
                    state['technical_review_version'] = previous_text
                    state['technical_issues'] = [{"agent": self.name, "type": "error", "content": str(e)}]
                    state['needs_human_review'] = False
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Error - using previous version")
        
        # INCREMENT INDEX
        state['current_agent_index'] = state.get('current_agent_index', 0) + 1
        return state

class LiteraryEditorAgent:
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
        self.name = "Literary Style & Excellence Editor"
        self.emoji = "✍️"
    
    async def polish(self, state: TranslationState) -> TranslationState:
        if state.get('literary_polish'):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Already complete - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        previous_text = (state.get('technical_review_version') or
                         state.get('tone_adjustment') or
                         state.get('cultural_adaptation') or
                         state.get('literal_translation') or "")
        
        if text_similarity(previous_text, state.get('literary_polish', '')) > 0.95:
            state['agent_notes'].append(f"{self.emoji} {self.name}: Minimal changes detected - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        attempt = 0
        max_attempts = 2
        
        while attempt < max_attempts:
            try:
                system_prompt = "You elevate prose to publication quality without altering meaning.\n\n" + LANGUAGE_GUARDRAIL
                
                user_prompt = f"""Polish this text to publication quality.

{LANGUAGE_GUARDRAIL}

**TEXT:**
{previous_text}

**AUDIENCE**: {state['target_audience']}

**YOUR TASKS:**
- Eliminate awkward phrasing
- Enhance word choice
- Optimize prose rhythm
- Strengthen imagery
- Maintain meaning

**OUTPUT**: Provide the polished text, then "LITERARY NOTES:" (notes must be in {state['target_language']})."""
                
                response = await self.llm.ainvoke([SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)])
                content = response.content
                
                polished, found_label, notes = split_notes(content, ["LITERARY NOTES:"])
                
                issues = []
                if notes:
                    issues.append({"agent": self.name, "type": "literary_notes", "content": notes})
                
                polished = await reinforce_language(self.llm, polished, state['target_language'])
                
                state['literary_polish'] = polished
                state['literary_issues'] = issues
                state['agent_notes'].append(f"{self.emoji} {self.name}: Literary polish completed")
                break
            
            except Exception as e:
                attempt += 1
                if attempt < max_attempts:
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Retry attempt {attempt} after error: {str(e)}")
                else:
                    st.error(f"Error in Literary step after {max_attempts} attempts: {str(e)}")
                    state['literary_polish'] = previous_text
                    state['literary_issues'] = [{"agent": self.name, "type": "error", "content": str(e)}]
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Error - using previous version")
        
        # INCREMENT INDEX
        state['current_agent_index'] = state.get('current_agent_index', 0) + 1
        return state

class QualityControlAgent:
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
        self.name = "Master Quality Synthesizer"
        self.emoji = "✅"
    
    async def finalize(self, state: TranslationState) -> TranslationState:
        if state.get('final_translation'):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Already complete - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        previous_text = (state.get('literary_polish') or
                         state.get('technical_review_version') or
                         state.get('tone_adjustment') or
                         state.get('cultural_adaptation') or
                         state.get('literal_translation') or "")
        
        if text_similarity(previous_text, state.get('final_translation', '')) > 0.95:
            state['agent_notes'].append(f"{self.emoji} {self.name}: Minimal changes detected - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        attempt = 0
        max_attempts = 2
        
        while attempt < max_attempts:
            try:
                system_prompt = "You integrate all contributions and output the final publication-ready text.\n\n" + LANGUAGE_GUARDRAIL
                
                all_issues = (
                    state.get('literal_issues', []) +
                    state.get('cultural_issues', []) +
                    state.get('tone_issues', []) +
                    state.get('technical_issues', []) +
                    state.get('literary_issues', [])
                )
                
                latest_version = (
                    state.get('literary_polish') or
                    state.get('technical_review_version') or
                    state.get('tone_adjustment') or
                    state.get('cultural_adaptation') or
                    state.get('literal_translation')
                )
                
                user_prompt = f"""Produce the final, publication-ready text.

{LANGUAGE_GUARDRAIL}

**LATEST VERSION:**
{latest_version}

**OUTPUT**: Provide ONLY the final text without meta-commentary, entirely in {state['target_language']}.
"""
                
                response = await self.llm.ainvoke([SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)])
                content = await reinforce_language(self.llm, response.content.strip(), state['target_language'])
                
                critical_passages = extract_critical_passages(content, all_issues)
                
                state['final_translation'] = content
                state['completed_at'] = datetime.now().isoformat()
                state['critical_passages'] = critical_passages
                state['agent_notes'].append(f"{self.emoji} {self.name}: Final output approved")
                
                # NEW: Calculate confidence scores
                confidence_scorer = ConfidenceScorer(self.llm)
                context = {
                    'source_lang': state['source_language'],
                    'target_lang': state['target_language'],
                    'entities': state.get('source_entities', [])
                }
                confidence = await confidence_scorer.score_translation(
                    state['source_text'],
                    content,
                    context
                )
                state['confidence_scores'] = confidence
                
                if 'enable_entity_tracking' in st.session_state and st.session_state.enable_entity_tracking:
                    entity_tracker = st.session_state.entity_tracker
                    state['translated_entities'] = entity_tracker.extract_entities(content)
                    
                    if state.get('source_entities'):
                        source_names = {e['name'].lower() for e in state['source_entities']}
                        translated_names = {e['name'].lower() for e in state['translated_entities']}
                        if source_names:
                            state['entity_preservation_rate'] = len(source_names & translated_names) / len(source_names)
                
                break
            
            except Exception as e:
                attempt += 1
                if attempt < max_attempts:
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Retry attempt {attempt} after error: {str(e)}")
                else:
                    st.error(f"Error in Finalize step after {max_attempts} attempts: {str(e)}")
                    state['final_translation'] = latest_version
                    state['completed_at'] = datetime.now().isoformat()
                    state['critical_passages'] = []
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Error - using latest version")
        
        # INCREMENT INDEX
        state['current_agent_index'] = state.get('current_agent_index', 0) + 1
        return state

class BERTScoreValidatorAgent:
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
        self.name = "Semantic Fidelity Validator"
        self.emoji = "🎯"
        self.target_score = 0.8
    
    async def validate_and_refine(self, state: TranslationState) -> TranslationState:
        """Validate BERTScore and refine if needed (same-language mode only)"""
        if state.get('bertscore_history') and state['bertscore_history'][-1]['f1'] >= self.target_score:
            state['agent_notes'].append(f"{self.emoji} {self.name}: Already validated - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        if not languages_equivalent(state['source_language'], state['target_language']):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Skipped (cross-language mode)")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        if not BERT_AVAILABLE:
            state['agent_notes'].append(f"{self.emoji} {self.name}: Skipped (bert-score not installed)")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        source_text = state['source_text']
        current_text = state['final_translation'] or ""
        
        if 'bertscore_attempts' not in state:
            state['bertscore_attempts'] = 0
        if 'bertscore_history' not in state:
            state['bertscore_history'] = []
        
        attempt = 0
        max_refinements = 20  # Safety cap to prevent infinite loops
        
        while attempt < max_refinements:
            state['bertscore_attempts'] += 1
            attempt += 1
            
            scores = compute_bertscore(current_text, source_text)
            
            if scores is None:
                state['agent_notes'].append(f"{self.emoji} {self.name}: BERTScore computation failed")
                break
            
            f1_score = scores['f1']
            state['bertscore_history'].append({
                'attempt': state['bertscore_attempts'],
                'f1': f1_score,
                'precision': scores['precision'],
                'recall': scores['recall']
            })
            
            if f1_score >= self.target_score:
                state['agent_notes'].append(
                    f"{self.emoji} {self.name}: ✅ BERTScore validated (F1={f1_score:.3f}) after {attempt} attempt(s)"
                )
                state['final_translation'] = current_text
                break
            
            if text_similarity(current_text, source_text) > 0.98:  # Near identical, no point refining
                state['agent_notes'].append(f"{self.emoji} {self.name}: Near-identical to source - accepting (F1={f1_score:.3f})")
                break
            
            state['agent_notes'].append(
                f"{self.emoji} {self.name}: 🔄 Refining (F1={f1_score:.3f} < {self.target_score}, attempt {attempt})"
            )
            
            system_prompt = (
                "You are a semantic fidelity specialist. Your ONLY goal is to increase semantic similarity "
                "to the source text while maintaining natural language quality.\n\n"
                + LANGUAGE_GUARDRAIL
            )
            
            user_prompt = f"""The current refined text has insufficient semantic similarity to the source (BERTScore F1: {f1_score:.3f}, target: {self.target_score}).

{LANGUAGE_GUARDRAIL}

**SOURCE TEXT (preserve its meaning):**
{source_text}

**CURRENT REFINED TEXT (needs closer alignment):**
{current_text}

**YOUR TASK:**
1. Identify where meaning has drifted from the source
2. Adjust ONLY those areas to restore semantic fidelity
3. DO NOT over-edit - preserve what is already good
4. Maintain natural, fluent language
5. Keep technical terms and proper nouns identical

**CRITICAL:** The goal is semantic similarity, not word-for-word copying. Preserve the SOURCE's meaning using natural language.

**OUTPUT:** Provide ONLY the refined text in {state['target_language']}, with no meta-commentary."""
            
            try:
                response = await self.llm.ainvoke([
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=user_prompt)
                ])
                current_text = await reinforce_language(self.llm, response.content.strip(), state['target_language'])
                
            except Exception as e:
                st.error(f"Error in BERTScore refinement: {str(e)}")
                state['agent_notes'].append(f"{self.emoji} {self.name}: Error during refinement - {str(e)}")
                break
        
        state['final_translation'] = current_text
        
        # INCREMENT INDEX
        state['current_agent_index'] = state.get('current_agent_index', 0) + 1
        return state

# =====================
# LLM Initialization Helper
# =====================
def initialize_llm(
    provider: Literal["openai", "anthropic"],
    model: str,
    api_key: str,
    temperature: float = 0.3
) -> BaseChatModel:
    if provider == "openai":
        if not OPENAI_AVAILABLE:
            raise ImportError("OpenAI not available. Install with: pip install langchain-openai")
        return ChatOpenAI(model=model, temperature=temperature, api_key=api_key, timeout=120)
    elif provider == "anthropic":
        if not ANTHROPIC_AVAILABLE:
            raise ImportError("Anthropic not available. Install with: pip install langchain-anthropic")
        return ChatAnthropic(model=model, temperature=temperature, api_key=api_key, timeout=120)
    else:
        raise ValueError(f"Unsupported provider: {provider}")

# =====================
# LangSmith Setup
# =====================
def setup_langsmith(api_key: Optional[str], project_name: str = "translation-pipeline"):
    if not api_key:
        return False
    if not LANGSMITH_AVAILABLE:
        st.warning("LangSmith not installed. Install with: pip install langsmith")
        return False
    
    try:
        os.environ["LANGCHAIN_TRACING_V2"] = "true"
        os.environ["LANGCHAIN_API_KEY"] = api_key
        os.environ["LANGCHAIN_PROJECT"] = project_name
        Client(api_key=api_key)
        return True
    except Exception as e:
        st.error(f"LangSmith setup failed: {str(e)}")
        return False

# =====================
# Graph Construction WITH DYNAMIC ROUTING
# =====================
def build_translation_graph(llm: BaseChatModel, cache: SemanticTranslationCache) -> StateGraph:
    planning_agent = PlanningAgent(llm)
    literal_agent = LiteralTranslationAgent(llm, cache)
    cultural_agent = CulturalAdaptationAgent(llm)
    tone_agent = ToneConsistencyAgent(llm)
    technical_agent = TechnicalReviewAgent(llm)
    literary_agent = LiteraryEditorAgent(llm)
    qc_agent = QualityControlAgent(llm)
    bertscore_agent = BERTScoreValidatorAgent(llm)
    
    workflow = StateGraph(TranslationState)
    
    workflow.add_node("planning", planning_agent.analyze_and_plan)
    workflow.add_node("literal_translator", literal_agent.translate)
    workflow.add_node("cultural_adapter", cultural_agent.adapt)
    workflow.add_node("tone_specialist", tone_agent.adjust_tone)
    workflow.add_node("technical_reviewer", technical_agent.review)
    workflow.add_node("literary_editor", literary_agent.polish)
    workflow.add_node("finalize", qc_agent.finalize)
    workflow.add_node("bertscore_validator", bertscore_agent.validate_and_refine)
    
    workflow.set_entry_point("planning")
    
    workflow.add_conditional_edges("planning", route_to_next_agent)
    workflow.add_conditional_edges("literal_translator", route_to_next_agent)
    workflow.add_conditional_edges("cultural_adapter", route_to_next_agent)
    workflow.add_conditional_edges("tone_specialist", route_to_next_agent)
    workflow.add_conditional_edges("technical_reviewer", route_to_next_agent)
    workflow.add_conditional_edges("literary_editor", route_to_next_agent)
    workflow.add_conditional_edges("finalize", route_to_next_agent)
    workflow.add_conditional_edges("bertscore_validator", route_to_next_agent)
    
    checkpointer = MemorySaver()
    return workflow.compile(checkpointer=checkpointer)

# =====================
# Streamlit UI - COMPLETE WITH NEW FEATURES
# =====================
async def main():
    st.set_page_config(
        page_title="Advanced Translation System",
        page_icon="🌐",
        layout="wide",
        initial_sidebar_state="expanded",
        menu_items=None
    )
    
    # Initialize session state
    if 'translation_state' not in st.session_state:
        st.session_state.translation_state = None
    if 'graph' not in st.session_state:
        st.session_state.graph = None
    if 'history' not in st.session_state:
        st.session_state.history = []
    if 'edited_translation' not in st.session_state:
        st.session_state.edited_translation = None
    if 'thread_id' not in st.session_state:
        st.session_state.thread_id = f"thread_{datetime.now().timestamp()}"
    
    # Cache initialization
    if 'translation_cache' not in st.session_state:
        st.session_state.translation_cache = SemanticTranslationCache()
    
    # Entity tracker initialization
    if 'entity_tracker' not in st.session_state:
        st.session_state.entity_tracker = EntitiesTracker()
    if 'enable_entity_tracking' not in st.session_state:
        st.session_state.enable_entity_tracking = False
    if 'enable_entity_awareness' not in st.session_state:
        st.session_state.enable_entity_awareness = False
    
    # Planning preferences
    if 'planning_enabled' not in st.session_state:
        st.session_state.planning_enabled = True
    if 'force_cultural_adapter' not in st.session_state:
        st.session_state.force_cultural_adapter = False
    if 'force_tone_specialist' not in st.session_state:
        st.session_state.force_tone_specialist = False
    if 'force_technical_reviewer' not in st.session_state:
        st.session_state.force_technical_reviewer = False
    if 'force_literary_editor' not in st.session_state:
        st.session_state.force_literary_editor = False
    if 'force_bertscore_validator' not in st.session_state:
        st.session_state.force_bertscore_validator = False
    
    # NEW: Alternative translations state
    if 'alternatives' not in st.session_state:
        st.session_state.alternatives = []
    
    # Custom CSS
    st.markdown("""
    <style>
        .agent-card {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            padding: 20px;
            border-radius: 10px;
            color: white;
            margin: 10px 0;
        }
        .critical-passage {
            background: #fff3cd;
            border-left: 4px solid #ffc107;
            padding: 10px;
            margin: 10px 0;
            border-radius: 4px;
        }
        .planning-card {
            background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
            padding: 15px;
            border-radius: 8px;
            color: white;
            margin: 10px 0;
        }
        .cache-card {
            background: linear-gradient(135deg, #00c6ff 0%, #0072ff 100%);
            padding: 15px;
            border-radius: 8px;
            color: white;
            margin: 10px 0;
        }
        .confidence-card {
            background: linear-gradient(135deg, #11998e 0%, #38ef7d 100%);
            padding: 15px;
            border-radius: 8px;
            color: white;
            margin: 10px 0;
        }
        .stTabs [data-baseweb="tab-list"] { gap: 24px; }
        .stTabs [data-baseweb="tab"] { padding: 10px 20px; }
    </style>
    """, unsafe_allow_html=True)
    
    st.title("🌐 Advanced Multi-Agent Translation System")
    st.caption("🎯 **NEW**: Confidence Scores • Diff Visualization • Alternative Translations • Smart Semantic Caching")
    
    # Sidebar
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        st.subheader("🤖 Model Provider")
        provider = st.radio(
            "Select Provider",
            ["openai", "anthropic"],
            format_func=lambda x: "OpenAI (GPT-4)" if x == "openai" else "Anthropic (Claude)"
        )
        
        st.subheader("🔑 API Keys")
        if provider == "openai":
            api_key = st.text_input("OpenAI API Key", type="password", help="Required for translation/refinement")
            model_options = ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo", "gpt-4", "gpt-3.5-turbo"]
        else:
            api_key = st.text_input("Anthropic API Key", type="password", help="Required for translation/refinement")
            model_options = ["claude-3-5-sonnet-20241022", "claude-3-opus-20240229", "claude-3-haiku-20240307"]
        
        st.subheader("📊 Monitoring (Optional)")
        enable_langsmith = st.checkbox("Enable LangSmith Tracing", help="Track and monitor agent interactions")
        langsmith_key = None
        langsmith_project = "translation-pipeline"
        
        if enable_langsmith:
            langsmith_key = st.text_input("LangSmith API Key", type="password")
            langsmith_project = st.text_input("LangSmith Project Name", value="translation-pipeline")
            if langsmith_key:
                if setup_langsmith(langsmith_key, langsmith_project):
                    st.success("✅ LangSmith enabled")
                else:
                    st.warning("⚠️ LangSmith setup failed")
        
        st.divider()
        
        st.subheader("🎛️ Model Settings")
        model = st.selectbox("Model", model_options, index=0)
        temperature = st.slider("Temperature", 0.0, 1.0, 0.3, help="Lower = more conservative, Higher = more creative")
        
        st.divider()
        
        # Cache Controls
        st.subheader("⚡ Smart Cache")
        
        cache_stats = st.session_state.translation_cache.get_stats()
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Entries", cache_stats['total_entries'])
        with col2:
            st.metric("Hit Rate", f"{cache_stats['hit_rate']:.1f}%")
        
        col3, col4 = st.columns(2)
        with col3:
            st.metric("Hits", cache_stats['hits'] + cache_stats['similar_hits'] + cache_stats['partial_hits'])
        with col4:
            st.metric("Misses", cache_stats['misses'])
        
        if cache_stats['total_entries'] > 0:
            with st.expander("📊 Cache Details"):
                st.write(f"**Exact hits:** {cache_stats['hits']}")
                st.write(f"**Similar hits:** {cache_stats['similar_hits']}")
                st.write(f"**Partial hits:** {cache_stats['partial_hits']}")
                st.write(f"**Sentence cache:** {cache_stats['sentence_entries']} entries")
                st.write(f"**Embeddings:** {cache_stats['embedding_entries']} stored")
        
        col_save, col_clear = st.columns(2)
        with col_save:
            if st.button("💾 Save", use_container_width=True):
                if st.session_state.translation_cache.save_cache():
                    st.success("✅ Saved")
        with col_clear:
            if st.button("🗑️ Clear", use_container_width=True):
                st.session_state.translation_cache.clear_cache()
                st.success("✅ Cleared")
                st.rerun()
        
        if cache_stats['total_entries'] > 0:
            cache_export = st.session_state.translation_cache.export_cache()
            st.download_button(
                "📥 Export Cache",
                json.dumps(cache_export, indent=2, default=str),
                file_name=f"translation_cache_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
                use_container_width=True
            )
        
        uploaded_cache = st.file_uploader("📤 Import Cache", type=['json'])
        if uploaded_cache:
            try:
                cache_data = json.load(uploaded_cache)
                st.session_state.translation_cache.import_cache(cache_data)
                st.success("✅ Cache imported")
                st.rerun()
            except Exception as e:
                st.error(f"Import failed: {str(e)}")
        
        if not SENTENCE_TRANSFORMERS_AVAILABLE:
            st.warning("⚠️ Install sentence-transformers for semantic caching: `pip install sentence-transformers`")
        
        st.divider()
        
        # Planning Controls
        st.subheader("📋 Intelligent Planning")
        st.session_state.planning_enabled = st.checkbox(
            "Enable Smart Planning",
            value=st.session_state.planning_enabled,
            help="Let AI select optimal agents for each task"
        )
        
        if not st.session_state.planning_enabled:
            st.info("⚠️ Planning disabled - all agents will run")
        else:
            st.success("✅ Smart planning active - optimizing workflow")
            
            with st.expander("🎯 Manual Overrides", expanded=False):
                st.caption("Force inclusion of specific agents:")
                st.session_state.force_cultural_adapter = st.checkbox(
                    "🌍 Force Cultural Adapter",
                    value=st.session_state.force_cultural_adapter
                )
                st.session_state.force_tone_specialist = st.checkbox(
                    "🎭 Force Tone Specialist",
                    value=st.session_state.force_tone_specialist
                )
                st.session_state.force_technical_reviewer = st.checkbox(
                    "🔬 Force Technical Reviewer",
                    value=st.session_state.force_technical_reviewer
                )
                st.session_state.force_literary_editor = st.checkbox(
                    "✍️ Force Literary Editor",
                    value=st.session_state.force_literary_editor
                )
                st.session_state.force_bertscore_validator = st.checkbox(
                    "🎯 Force BERTScore Validator",
                    value=st.session_state.force_bertscore_validator,
                    help="Only applies to same-language refinement"
                )
        
        st.divider()
        
        st.subheader("🌍 Translation Settings")
        source_lang = st.selectbox(
            "Source Language",
            [
                "Ukrainian","Russian","Polish","German","French","Spanish","Italian","Portuguese",
                "English (US)","English (UK)","Romanian","Czech","Slovak","Bulgarian","Dutch","Other",
            ],
            help="Language of the source text"
        )
        
        target_lang = st.selectbox(
            "Target Language",
            ["English (US)","English (UK)","Spanish","French","German","Other"],
            help="Language to translate into"
        )
        
        same_lang_note = languages_equivalent(source_lang, target_lang)
        if same_lang_note:
            st.info("Same-language detected: running in **Refine mode** (copy-edit without changing meaning).")
        
        audience = st.selectbox(
            "Target Audience",
            [
                "General wellness readers","Academic/Technical audience","Business professionals",
                "Literary fiction readers","Young adults","Healthcare professionals","Scientific community"
            ],
            help="Who will read this output?"
        )
        
        genre = st.selectbox(
            "Content Genre",
            [
                "Wellness/Self-help","Literary Fiction","Academic/Technical",
                "Business/Professional","Scientific/Medical","General Non-fiction"
            ],
            help="Type of content"
        )
        
        st.divider()
        show_charts = st.checkbox("Show analytics visualizations", value=True)
        
        # Entity Tracking Options
        st.divider()
        st.subheader("🎯 Entity Tracking")
        st.session_state.enable_entity_tracking = st.checkbox(
            "Enable entity tracking",
            value=st.session_state.enable_entity_tracking,
            help="Track entities (people, places, organizations) through translation"
        )
        
        if st.session_state.enable_entity_tracking:
            st.session_state.enable_entity_awareness = st.checkbox(
                "Entity-aware agents",
                value=st.session_state.enable_entity_awareness,
                help="Give translation agents awareness of important entities"
            )
            
            uploaded_glossary = st.file_uploader(
                "Upload Entity Glossary",
                type=['json', 'csv', 'txt'],
                help="Upload a glossary of important terms to track"
            )
            
            if uploaded_glossary:
                if st.session_state.entity_tracker.upload_glossary(uploaded_glossary):
                    st.success(f"✅ Glossary loaded")
            
            with st.expander("➕ Quick Add Term"):
                new_term = st.text_input("Term name")
                term_type = st.selectbox("Type", ["person", "location", "organization", "date", "custom"])
                term_desc = st.text_input("Description (optional)")
                if st.button("Add to Glossary"):
                    if new_term:
                        st.session_state.entity_glossary[new_term] = {
                            'type': term_type,
                            'description': term_desc,
                            'aliases': []
                        }
                        st.success(f"Added: {new_term}")
                        st.rerun()
            
            st.metric("Glossary Terms", len(st.session_state.entity_glossary))
    
    # Main content
    st.header(f"📝 Interface · {mode_phrase(source_lang, target_lang)}")
    
    with st.expander("📚 Learn About Our Translation Agents", expanded=False):
        st.markdown("### The Eight-Agent Pipeline (with Intelligent Planning)")
        for agent_key, agent_info in AGENT_DESCRIPTIONS.items():
            st.markdown(f"""
            <div class="agent-card">
                <h3>{agent_info['emoji']} {agent_info['name']}</h3>
                <h4>Role: {agent_info['role']}</h4>
            </div>
            """, unsafe_allow_html=True)
            st.markdown(agent_info['description'])
            st.markdown("**Areas of Expertise:**")
            cols = st.columns(len(agent_info['expertise']))
            for idx, ex in enumerate(agent_info['expertise']):
                cols[idx].info(ex)
            st.divider()
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📄 Source Text")
        uploaded_file = st.file_uploader("Upload a file (optional)", type=['txt','docx','md'])
        
        if uploaded_file:
            file_content = read_uploaded_file(uploaded_file)
            if file_content:
                st.success(f"✅ File loaded: {uploaded_file.name}")
                source_text = st.text_area("Edit source text if needed", value=file_content, height=400)
            else:
                source_text = st.text_area("Enter text", height=400, placeholder="Paste text manually...")
        else:
            source_text = st.text_area(
                "Enter text",
                height=400,
                placeholder="Paste your source text here…",
            )
        
        col_a, col_b = st.columns([3, 1])
        with col_a:
            translate_button = st.button(
                "🚀 Start Pipeline",
                type="primary",
                disabled=not api_key,
                use_container_width=True
            )
        with col_b:
            if st.button("🗑️ Clear", use_container_width=True):
                st.session_state.translation_state = None
                st.session_state.edited_translation = None
                st.session_state.alternatives = []
                st.session_state.thread_id = f"thread_{datetime.now().timestamp()}"
                st.rerun()
        
        if not api_key:
            st.warning("⚠️ Enter your API key in the sidebar to begin")
        
        if source_text:
            st.caption(f"📊 {len(source_text)} characters | ~{len(source_text.split())} words")
    
    with col2:
        st.subheader("🎯 Result")
        
        if translate_button and source_text:
            try:
                status_container = st.empty()
                progress_container = st.empty()
                
                with status_container.container():
                    st.info("🔄 Initializing pipeline...")
                
                llm = initialize_llm(provider=provider, model=model, api_key=api_key, temperature=temperature)
                cache = st.session_state.translation_cache
                graph = build_translation_graph(llm, cache)
                st.session_state.graph = graph
                
                current_time = datetime.now().isoformat()
                
                initial_state = {
                    "source_text": source_text,
                    "source_language": source_lang,
                    "target_language": target_lang,
                    "target_audience": audience,
                    "genre": genre,
                    "literal_translation": "",
                    "cultural_adaptation": "",
                    "tone_adjustment": "",
                    "technical_review_version": "",
                    "literary_polish": "",
                    "final_translation": "",
                    "literal_issues": [],
                    "cultural_issues": [],
                    "tone_issues": [],
                    "technical_issues": [],
                    "literary_issues": [],
                    "critical_passages": [],
                    "agent_notes": [],
                    "agent_decisions": [],
                    "human_feedback": None,
                    "revision_count": 0,
                    "needs_human_review": False,
                    "bertscore_attempts": 0,
                    "bertscore_history": [],
                    "started_at": current_time,
                    "completed_at": None,
                    "agent_plan": [],
                    "agent_plan_reasoning": {},
                    "planning_analysis": {},
                    "skipped_agents": [],
                    "estimated_complexity": "unknown",
                    "current_agent_index": 0,
                    "planning_enabled": st.session_state.planning_enabled,
                    "cache_hit": None,
                    "cache_speedup": None,
                    "confidence_scores": None
                }
                
                if st.session_state.enable_entity_tracking:
                    entity_tracker = st.session_state.entity_tracker
                    source_entities = entity_tracker.extract_entities(source_text)
                    initial_state['source_entities'] = source_entities
                    initial_state['translated_entities'] = []
                    initial_state['entity_preservation_rate'] = 0.0
                    
                    if source_entities:
                        st.info(f"🎯 Tracking {len(source_entities)} entities through translation")
                else:
                    initial_state['source_entities'] = None
                    initial_state['translated_entities'] = None
                    initial_state['entity_preservation_rate'] = None
                
                with status_container.container():
                    st.info("🔄 Planning workflow...")
                with progress_container.container():
                    st.progress(0.05)
                
                with status_container.container():
                    st.info("🔄 Running agents...")
                
                config = {"configurable": {"thread_id": st.session_state.thread_id}}
                result = None
                
                try:
                    result = await graph.ainvoke(initial_state, config)
                except Exception as inner_e:
                    st.warning("Partial failure detected - attempting resume from checkpoint...")
                    state_snapshot = graph.checkpointer.get_tuple(config)
                    if state_snapshot:
                        result = await graph.ainvoke(None, config)
                    else:
                        raise inner_e
                
                status_container.empty()
                progress_container.empty()
                
                st.session_state.translation_state = result
                st.session_state.edited_translation = result.get('final_translation', '') if result else ''
                st.session_state.alternatives = []  # Reset alternatives
                
                st.session_state.history.append({
                    "timestamp": datetime.now().isoformat(),
                    "source": source_text[:100] + "...",
                    "target_lang": target_lang,
                    "model": model,
                    "result": result
                })
                
                # Show cache impact
                if result.get('cache_hit'):
                    cache_type = result['cache_hit']
                    speedup = result.get('cache_speedup', 'unknown')
                    st.success(f"✅ Pipeline completed! ⚡ CACHE HIT ({cache_type}) - {speedup} speedup")
                else:
                    st.success("✅ Pipeline completed!")
                
                st.rerun()
            
            except Exception as e:
                st.error(f"❌ Error during processing: {str(e)}")
                st.code(traceback.format_exc())
        
        if st.session_state.translation_state:
            state = st.session_state.translation_state
            
            # Display cache hit if applicable
            if state.get('cache_hit'):
                st.markdown(f"""
                <div class="cache-card">
                    <h3>⚡ Cache Hit: {state['cache_hit'].upper()}</h3>
                    <p>Speedup: {state.get('cache_speedup', 'significant')}</p>
                </div>
                """, unsafe_allow_html=True)
            
            # NEW: Display confidence scores
            if state.get('confidence_scores'):
                confidence = state['confidence_scores']
                st.markdown(f"""
                <div class="confidence-card">
                    <h3>🎯 Translation Confidence: {confidence['overall']:.0%}</h3>
                    <p>Fluency: {confidence.get('fluency', 0.5):.0%} | Fidelity: {confidence.get('semantic_fidelity', 0.5):.0%} | Terminology: {confidence.get('terminology', 1.0):.0%}</p>
                </div>
                """, unsafe_allow_html=True)
            
            # Display planning information
            if state.get('agent_plan'):
                st.markdown("### 📋 Execution Plan")
                
                plan = state.get('agent_plan', [])
                skipped = state.get('skipped_agents', [])
                complexity = state.get('estimated_complexity', 'unknown')
                
                agent_names = {
                    'literal_translator': '🔤 Baseline',
                    'cultural_adapter': '🌍 Cultural',
                    'tone_specialist': '🎭 Tone',
                    'technical_reviewer': '🔬 Technical',
                    'literary_editor': '✍️ Literary',
                    'finalize': '✅ Finalize',
                    'bertscore_validator': '🎯 BERTScore'
                }
                
                col_info1, col_info2, col_info3 = st.columns(3)
                with col_info1:
                    st.metric("Agents Used", f"{len(plan)}/7")
                with col_info2:
                    savings = int((1 - len(plan) / 7) * 100)
                    st.metric("Efficiency Gain", f"{savings}%")
                with col_info3:
                    st.metric("Complexity", complexity.title())
                
                route_display = " → ".join([agent_names.get(a, a) for a in plan])
                st.info(f"**Route:** {route_display}")
                
                if skipped:
                    skipped_display = ", ".join([agent_names.get(a, a) for a in skipped])
                    st.caption(f"⏭️ **Skipped:** {skipped_display}")
            
            st.divider()
            st.markdown("### 📖 Final Output")
            display_text = st.session_state.edited_translation or state.get('final_translation', 'No output yet')
            st.text_area("Publication-Ready Text", display_text, height=400, key="final_display", label_visibility="visible")
            
            # Downloads
            st.markdown("### 📥 Download Options")
            col_x, col_y, col_z, col_w = st.columns(4)
            
            final_text = st.session_state.edited_translation or state.get('final_translation', '')
            
            with col_x:
                st.download_button(
                    "📄 Text (.txt)", final_text,
                    file_name=f"translation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain", use_container_width=True
                )
            
            with col_y:
                md_content = create_markdown_file(final_text, "Translation")
                st.download_button(
                    "📝 Markdown (.md)", md_content,
                    file_name=f"translation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                    mime="text/markdown", use_container_width=True
                )
            
            with col_z:
                if DOCX_AVAILABLE:
                    docx_file = create_docx_file(final_text, "Translation")
                    if docx_file:
                        st.download_button(
                            "📘 Word (.docx)", docx_file,
                            file_name=f"translation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                else:
                    st.button("📘 Word (.docx)", disabled=True, help="Install python-docx: pip install python-docx", use_container_width=True)
            
            with col_w:
                st.download_button(
                    "📊 Report (.json)", json.dumps(state, indent=2, default=str),
                    file_name=f"translation_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json", use_container_width=True
                )
    
    # Detailed Analysis Tabs
    if st.session_state.translation_state:
        st.divider()
        st.header("🔍 Detailed Analysis")
        
        state = st.session_state.translation_state
        
        tab_names = [
            "🎯 Confidence","🔀 Diff Viewer","🎲 Alternatives",
            "⚡ Cache","📋 Planning","✏️ Edit & Review","🔄 Agent Workflow",
            "📊 All Versions","⚠️ Issues & Feedback","📈 Analytics","📚 History"
        ]
        
        if st.session_state.enable_entity_tracking:
            tab_names.append("🎯 Entities")
        
        tabs = st.tabs(tab_names)
        
        # Unpack tabs
        if len(tabs) == 12:
            (tab_confidence, tab_diff, tab_alternatives, tab_cache, tab_planning, 
             tab_edit, tab_workflow, tab_versions, tab_issues, tab_analytics, 
             tab_history, tab_entities) = tabs
        else:
            (tab_confidence, tab_diff, tab_alternatives, tab_cache, tab_planning, 
             tab_edit, tab_workflow, tab_versions, tab_issues, tab_analytics, 
             tab_history) = tabs
            tab_entities = None
        
        # NEW: Confidence Tab
        with tab_confidence:
            st.subheader("🎯 Translation Confidence Scores")
            
            if state.get('confidence_scores'):
                confidence = state['confidence_scores']
                
                # Overall score display
                overall = confidence['overall']
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("Overall", f"{overall:.1%}", 
                             help="Weighted average of all metrics")
                with col2:
                    st.metric("Fluency", f"{confidence.get('fluency', 0.5):.1%}",
                             help="Natural language quality")
                with col3:
                    st.metric("Semantic Fidelity", f"{confidence.get('semantic_fidelity', 0.5):.1%}",
                             help="Meaning preservation (BERTScore)")
                with col4:
                    st.metric("Terminology", f"{confidence.get('terminology', 1.0):.1%}",
                             help="Key term preservation")
                
                # Visual gauge
                if PLOTLY_AVAILABLE:
                    st.divider()
                    st.markdown("### 📊 Confidence Visualization")
                    
                    # Gauge chart
                    fig = go.Figure(go.Indicator(
                        mode = "gauge+number",
                        value = overall * 100,
                        domain = {'x': [0, 1], 'y': [0, 1]},
                        title = {'text': "Overall Confidence"},
                        gauge = {
                            'axis': {'range': [None, 100]},
                            'bar': {'color': "darkblue"},
                            'steps': [
                                {'range': [0, 50], 'color': "lightgray"},
                                {'range': [50, 75], 'color': "gray"},
                                {'range': [75, 100], 'color': "lightgreen"}
                            ],
                            'threshold': {
                                'line': {'color': "red", 'width': 4},
                                'thickness': 0.75,
                                'value': 80
                            }
                        }
                    ))
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Bar chart of components
                    st.markdown("### 📈 Component Breakdown")
                    metrics_data = {
                        'Metric': ['Fluency', 'Semantic Fidelity', 'Terminology', 'Length Ratio'],
                        'Score': [
                            confidence.get('fluency', 0.5) * 100,
                            confidence.get('semantic_fidelity', 0.5) * 100,
                            confidence.get('terminology', 1.0) * 100,
                            confidence.get('length_ratio', 1.0) * 100
                        ]
                    }
                    
                    fig2 = go.Figure(data=[
                        go.Bar(
                            x=metrics_data['Metric'],
                            y=metrics_data['Score'],
                            marker_color=['#3b82f6', '#10b981', '#f59e0b', '#8b5cf6']
                        )
                    ])
                    fig2.update_layout(
                        title='Confidence Metrics',
                        yaxis_title='Score (%)',
                        yaxis_range=[0, 100],
                        showlegend=False
                    )
                    st.plotly_chart(fig2, use_container_width=True)
                
                # Interpretation guide
                st.divider()
                st.markdown("### 📖 Score Interpretation")
                
                if overall >= 0.85:
                    st.success("✅ **Excellent** - High confidence translation")
                elif overall >= 0.7:
                    st.info("ℹ️ **Good** - Solid translation with minor areas for review")
                elif overall >= 0.5:
                    st.warning("⚠️ **Fair** - Consider reviewing specific sections")
                else:
                    st.error("❌ **Poor** - Significant revision recommended")
                
                with st.expander("📚 Understanding the Metrics"):
                    st.markdown("""
                    **Overall Confidence:** Weighted combination of all metrics
                    - 35% Semantic Fidelity
                    - 35% Fluency
                    - 15% Terminology
                    - 15% Length Ratio
                    
                    **Fluency:** How natural and well-written the translation reads
                    
                    **Semantic Fidelity:** How well the meaning is preserved (uses BERTScore for same-language)
                    
                    **Terminology:** Percentage of key entities/terms preserved
                    
                    **Length Ratio:** How appropriate the length is (ideal: 0.8-1.2x source)
                    """)
            else:
                st.info("Confidence scores not available for this translation")
        
        # NEW: Diff Viewer Tab
        with tab_diff:
            st.subheader("🔀 Visual Diff Between Agent Versions")
            
            versions = [
                ("Baseline", state.get('literal_translation', '')),
                ("Cultural", state.get('cultural_adaptation', '')),
                ("Tone", state.get('tone_adjustment', '')),
                ("Technical", state.get('technical_review_version', '')),
                ("Literary", state.get('literary_polish', '')),
                ("Final", state.get('final_translation', ''))
            ]
            
            # Filter out empty versions
            available_versions = [(name, text) for name, text in versions if text]
            
            if len(available_versions) >= 2:
                st.markdown("### 📊 Change Rate Summary")
                
                # Calculate change rates between consecutive versions
                change_data = []
                for i in range(len(available_versions) - 1):
                    name1, text1 = available_versions[i]
                    name2, text2 = available_versions[i + 1]
                    change_rate = calculate_change_rate(text1, text2)
                    change_data.append({
                        'Transition': f"{name1} → {name2}",
                        'Change Rate': change_rate
                    })
                
                if change_data:
                    df_changes = pd.DataFrame(change_data)
                    
                    if PLOTLY_AVAILABLE:
                        fig = go.Figure(data=[
                            go.Bar(
                                x=df_changes['Transition'],
                                y=df_changes['Change Rate'],
                                marker_color='#3b82f6'
                            )
                        ])
                        fig.update_layout(
                            title='Change Rate Between Agent Versions',
                            yaxis_title='Change Rate (%)',
                            xaxis_title='Agent Transition',
                            showlegend=False
                        )
                        st.plotly_chart(fig, use_container_width=True)
                    else:
                        st.dataframe(df_changes)
                
                st.divider()
                st.markdown("### 🔍 Detailed Comparison")
                
                # Version selector
                col1, col2 = st.columns(2)
                with col1:
                    version1_idx = st.selectbox(
                        "Compare from:",
                        range(len(available_versions)),
                        format_func=lambda i: available_versions[i][0]
                    )
                with col2:
                    version2_idx = st.selectbox(
                        "Compare to:",
                        range(len(available_versions)),
                        index=min(version1_idx + 1, len(available_versions) - 1),
                        format_func=lambda i: available_versions[i][0]
                    )
                
                if version1_idx != version2_idx:
                    name1, text1 = available_versions[version1_idx]
                    name2, text2 = available_versions[version2_idx]
                    
                    change_rate = calculate_change_rate(text1, text2)
                    st.info(f"**Change Rate:** {change_rate:.1f}% of the text was modified")
                    
                    diff_html = create_diff_visualization(text1, text2, name1, name2)
                    st.markdown(diff_html, unsafe_allow_html=True)
                else:
                    st.warning("Please select two different versions to compare")
            else:
                st.info("Need at least 2 versions to show diff")
        
        # NEW: Alternatives Tab
        with tab_alternatives:
            st.subheader("🎲 Alternative Translation Variants")
            
            if not st.session_state.alternatives:
                st.markdown("""
                Generate alternative translations using different strategies:
                - **Conservative**: More literal and faithful
                - **Balanced**: Mix of literal and natural
                - **Creative**: More idiomatic and natural
                """)
                
                num_alts = st.slider("Number of alternatives", 2, 3, 3)
                
                if st.button("🎲 Generate Alternatives", type="primary"):
                    with st.spinner("Generating alternative translations..."):
                        try:
                            llm = initialize_llm(provider=provider, model=model, api_key=api_key, temperature=temperature)
                            generator = AlternativeTranslationGenerator(llm)
                            alternatives = await generator.generate_alternatives(state, num_alts)
                            st.session_state.alternatives = alternatives
                            st.success(f"✅ Generated {len(alternatives)} alternatives")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error generating alternatives: {str(e)}")
            else:
                st.success(f"✅ {len(st.session_state.alternatives)} alternatives generated")
                
                # Display current translation for comparison
                with st.expander("📖 Current Translation (for comparison)", expanded=False):
                    st.text_area(
                        "Current",
                        st.session_state.edited_translation,
                        height=150,
                        key="current_for_comparison",
                        disabled=True
                    )
                
                st.divider()
                
                # Display alternatives
                for idx, alt in enumerate(st.session_state.alternatives, 1):
                    if alt.get('error'):
                        st.error(f"❌ Alternative {idx} ({alt['strategy']}): Failed to generate")
                        continue
                    
                    st.markdown(f"### 🎯 Alternative {idx}: {alt['strategy'].title()}")
                    st.caption(f"📝 {alt['description']} (temperature: {alt['temperature']})")
                    
                    st.text_area(
                        f"Alt {idx}",
                        alt['translation'],
                        height=150,
                        key=f"alt_display_{idx}",
                        label_visibility="collapsed"
                    )
                    
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        if st.button(f"✅ Use This Version", key=f"use_alt_{idx}"):
                            st.session_state.edited_translation = alt['translation']
                            st.success("Alternative selected! Scroll up to see it in the Result section.")
                            st.rerun()
                    
                    with col2:
                        if st.button(f"⚖️ Compare with Current", key=f"compare_alt_{idx}"):
                            st.markdown("#### Comparison")
                            current = st.session_state.edited_translation
                            change_rate = calculate_change_rate(current, alt['translation'])
                            st.info(f"Change rate: {change_rate:.1f}%")
                            diff_html = create_diff_visualization(
                                current,
                                alt['translation'],
                                "Current",
                                f"Alternative {idx}"
                            )
                            st.markdown(diff_html, unsafe_allow_html=True)
                    
                    with col3:
                        # Word count comparison
                        current_words = len(st.session_state.edited_translation.split())
                        alt_words = len(alt['translation'].split())
                        word_diff = alt_words - current_words
                        st.metric("Words", alt_words, delta=word_diff)
                    
                    st.divider()
                
                if st.button("🔄 Generate New Alternatives"):
                    st.session_state.alternatives = []
                    st.rerun()
        
        # Cache Tab
        with tab_cache:
            st.subheader("⚡ Smart Cache Performance")
            
            cache_stats = st.session_state.translation_cache.get_stats()
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Total Entries", cache_stats['total_entries'])
            with col2:
                st.metric("Hit Rate", f"{cache_stats['hit_rate']:.1f}%")
            with col3:
                st.metric("Total Hits", cache_stats['hits'] + cache_stats['similar_hits'] + cache_stats['partial_hits'])
            with col4:
                st.metric("Misses", cache_stats['misses'])
            
            if state.get('cache_hit'):
                st.success(f"✅ This translation used cache: **{state['cache_hit'].upper()}** hit (speedup: {state.get('cache_speedup', 'significant')})")
            else:
                st.info("ℹ️ This translation was processed from scratch (no cache hit)")
            
            st.divider()
            
            st.markdown("### 📊 Cache Breakdown")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Exact Hits", cache_stats['hits'], help="100% match - instant result")
            with col2:
                st.metric("Similar Hits", cache_stats['similar_hits'], help="High similarity - light refinement")
            with col3:
                st.metric("Partial Hits", cache_stats['partial_hits'], help="Some sentences cached")
            
            if cache_stats['total_entries'] > 0:
                st.markdown("### 📦 Cache Contents")
                st.write(f"**Sentence-level cache:** {cache_stats['sentence_entries']} entries")
                st.write(f"**Embeddings stored:** {cache_stats['embedding_entries']}")
                
                if SENTENCE_TRANSFORMERS_AVAILABLE:
                    st.success("✅ Semantic similarity matching enabled")
                else:
                    st.warning("⚠️ Install sentence-transformers for semantic caching")
        
        # Planning Tab
        with tab_planning:
            st.subheader("📋 Planning Analysis & Decisions")
            
            if state.get('agent_plan'):
                reasoning = state.get('agent_plan_reasoning', {})
                analysis = state.get('planning_analysis', {})
                
                st.markdown("#### 📊 Text Analysis")
                if analysis:
                    cols = st.columns(3)
                    if 'length' in analysis:
                        cols[0].metric("Length", f"{analysis['length']} words")
                    if 'complexity_score' in analysis:
                        cols[1].metric("Complexity", f"{analysis.get('complexity_score', 0):.1f}/10")
                    if 'technical_terms_count' in analysis:
                        cols[2].metric("Technical Terms", analysis.get('technical_terms_count', 0))
                    
                    cols2 = st.columns(3)
                    if 'cultural_references_count' in analysis:
                        cols2[0].metric("Cultural Refs", analysis.get('cultural_references_count', 0))
                    if 'tone_consistency' in analysis:
                        cols2[1].metric("Tone", analysis.get('tone_consistency', 'N/A'))
                    if 'literary_elements' in analysis:
                        cols2[2].metric("Literary", analysis.get('literary_elements', 'N/A'))
                
                st.divider()
                st.markdown("#### 🎯 Agent Selection Reasoning")
                
                for agent_key in ['literal_translator', 'cultural_adapter', 'tone_specialist',
                                 'technical_reviewer', 'literary_editor', 'finalize', 'bertscore_validator']:
                    if agent_key in reasoning:
                        reason = reasoning[agent_key]
                        if reason.startswith('✅'):
                            st.success(f"**{agent_key.replace('_', ' ').title()}:** {reason}")
                        else:
                            st.info(f"**{agent_key.replace('_', ' ').title()}:** {reason}")
            else:
                st.info("Planning information not available")
        
        # Edit & Review Tab
        with tab_edit:
            st.subheader("✏️ Edit & Review")
            edited_text = st.text_area(
                "Edit the output as needed",
                value=st.session_state.edited_translation or state.get('final_translation', ''),
                height=300, key="edit_translation_area", label_visibility="visible"
            )
            
            if edited_text != st.session_state.edited_translation:
                st.session_state.edited_translation = edited_text
                st.success("✅ Edits saved")
            
            col_reset, col_apply = st.columns([1, 1])
            with col_reset:
                if st.button("↺ Reset to Original", use_container_width=True):
                    st.session_state.edited_translation = state.get('final_translation', '')
                    st.rerun()
            with col_apply:
                if st.button("💾 Save Final Version", type="primary", use_container_width=True):
                    state['final_translation'] = st.session_state.edited_translation
                    st.success("✅ Final version saved!")
            
            st.divider()
            st.markdown("#### 🚩 Critical Passages")
            critical_passages = state.get('critical_passages', [])
            
            if critical_passages:
                st.info(f"Found {len(critical_passages)} passage(s) that may need attention")
                for idx, p in enumerate(critical_passages, 1):
                    st.markdown(f"""
                    <div class="critical-passage">
                        <strong>Passage {idx}</strong> - Flagged by: {p.get('agent','Unknown')}
                    </div>
                    """, unsafe_allow_html=True)
                    st.markdown(f"**Text:** {p.get('passage','N/A')}")
                    st.markdown(f"**Issue:** {p.get('issue','No details available')}")
                    st.markdown(f"**Type:** `{p.get('type','general')}`")
                    
                    with st.expander(f"✏️ Edit Passage {idx}"):
                        new_p = st.text_area(f"Edit passage {idx}", value=p.get('passage',''), key=f"edit_passage_{idx}", height=100)
                        if st.button(f"Apply Edit to Passage {idx}", key=f"apply_{idx}"):
                            old = p.get('passage','')
                            if old in st.session_state.edited_translation:
                                st.session_state.edited_translation = st.session_state.edited_translation.replace(old, new_p)
                                st.success(f"✅ Passage {idx} updated")
                                st.rerun()
                    st.divider()
            else:
                st.success("✅ No critical passages flagged")
            
            review_notes = st.text_area("📝 Review Notes", placeholder="Enter notes/feedback…", height=150, key="review_notes")
            if st.button("💾 Save Review Notes"):
                if 'review_notes' not in state:
                    state['review_notes'] = []
                state['review_notes'].append({'timestamp': datetime.now().isoformat(), 'notes': review_notes})
                st.success("✅ Review notes saved")
        
        # Agent Workflow Tab
        with tab_workflow:
            st.subheader("🔄 Agent Workflow Progress")
            unique_notes = []
            seen = set()
            for note in state.get('agent_notes', []):
                if note not in seen:
                    unique_notes.append(note)
                    seen.add(note)
            
            for note in unique_notes:
                if "CACHE HIT" in note:
                    st.success(f"⚡ {note}")
                else:
                    st.success(f"✓ {note}")
            
            st.divider()
            
            if state.get('started_at') and state.get('completed_at'):
                try:
                    start = datetime.fromisoformat(state['started_at'])
                    end = datetime.fromisoformat(state['completed_at'])
                    duration = (end - start).total_seconds()
                    st.metric("Total Processing Time", f"{duration:.1f} seconds")
                except:
                    st.info("Processing time not available")
        
        # All Versions Tab
        with tab_versions:
            st.subheader("📊 All Versions")
            versions = [
                ("1️⃣ Baseline", state.get('literal_translation','')),
                ("2️⃣ Cultural/Register", state.get('cultural_adaptation','')),
                ("3️⃣ Tone", state.get('tone_adjustment','')),
                ("4️⃣ Technical", state.get('technical_review_version','')),
                ("5️⃣ Literary", state.get('literary_polish','')),
                ("6️⃣ Final", state.get('final_translation','')),
            ]
            
            for title, content in versions:
                if content:
                    with st.expander(title, expanded=False):
                        st.text_area(f"{title} content", content, height=200, key=f"version_{title}", label_visibility="collapsed")
        
        # Issues & Feedback Tab
        with tab_issues:
            st.subheader("⚠️ Issues & Feedback")
            all_issues = (
                state.get('literal_issues', []) +
                state.get('cultural_issues', []) +
                state.get('tone_issues', []) +
                state.get('technical_issues', []) +
                state.get('literary_issues', [])
            )
            
            if all_issues:
                for issue in all_issues:
                    with st.expander(f"{issue.get('agent','Unknown')} - {issue.get('type','general')}"):
                        st.markdown(issue.get('content',''))
            else:
                st.info("✅ No issues flagged")
            
            if state.get('needs_human_review'):
                st.warning("⚠️ Flagged for human review")
                feedback = st.text_area("Provide feedback for revision", placeholder="Describe needed changes…")
                if st.button("Submit Feedback"):
                    st.info("Feedback captured (extend workflow to loop back if desired).")
        
        # Analytics Tab
        with tab_analytics:
            st.subheader("📈 Analytics")
            col_a, col_b, col_c = st.columns(3)
            
            with col_a:
                source_words = len(state['source_text'].split())
                final_words = len(state.get('final_translation','').split())
                st.metric("Word Count", f"{final_words}", delta=f"{final_words - source_words}")
            with col_b:
                st.metric("Agents Involved", len(state.get('agent_notes', [])))
            with col_c:
                total_issues = sum(len(state.get(k, [])) for k in ['literal_issues','cultural_issues','tone_issues','technical_issues','literary_issues'])
                st.metric("Total Issues Addressed", total_issues)
            
            if show_charts:
                st.divider()
                with st.expander("📊 Word Count Overview", expanded=True):
                    try:
                        df_counts = pd.DataFrame(
                            {"Text": ["Source", "Final"], "Words": [source_words, final_words]}
                        ).set_index("Text")
                        st.bar_chart(df_counts)
                    except Exception:
                        st.warning("Could not render word count chart.")
                
                st.divider()
                with st.expander("📏 Sentence Length Distribution (words per sentence)", expanded=False):
                    src_lens = sentence_lengths(state.get('source_text', ''))
                    fin_lens = sentence_lengths(state.get('final_translation', ''))
                    max_len_for_bins = max([0] + src_lens + fin_lens)
                    bins = nonzero_bins(max_len_for_bins)
                    
                    cols = st.columns(2)
                    with cols[0]:
                        st.caption("Source")
                        try:
                            fig1, ax1 = plt.subplots()
                            ax1.hist(src_lens, bins=bins)
                            ax1.set_xlabel("Words per sentence")
                            ax1.set_ylabel("Frequency")
                            ax1.set_title("Source")
                            st.pyplot(fig1)
                        except Exception:
                            st.warning("Could not render source histogram.")
                    with cols[1]:
                        st.caption("Final")
                        try:
                            fig2, ax2 = plt.subplots()
                            ax2.hist(fin_lens, bins=bins)
                            ax2.set_xlabel("Words per sentence")
                            ax2.set_ylabel("Frequency")
                            ax2.set_title("Final")
                            st.pyplot(fig2)
                        except Exception:
                            st.warning("Could not render final histogram.")
                
                st.divider()
                with st.expander("☁️ Word Clouds (Before / After / Difference)", expanded=True):
                    if not WORDCLOUD_AVAILABLE:
                        st.info("Install `wordcloud` to enable this: `pip install wordcloud`")
                    else:
                        sw = set(STOPWORDS) | {"—", "–", "'", """, """, "…"}
                        src_text = state.get('source_text', '')
                        fin_text = state.get('final_translation', '')
                        
                        src_freq = word_frequencies(src_text, stopwords=sw)
                        fin_freq = word_frequencies(fin_text, stopwords=sw)
                        
                        diff_freq = {}
                        for w, f_cnt in fin_freq.items():
                            s_cnt = src_freq.get(w, 0)
                            delta = f_cnt - s_cnt
                            if delta > 0:
                                diff_freq[w] = delta
                        
                        st.caption("Before (Source)")
                        try:
                            render_wordcloud_from_freq(src_freq, "Source Word Cloud")
                        except Exception:
                            st.warning("Could not render source word cloud.")
                        
                        st.caption("After (Final)")
                        try:
                            render_wordcloud_from_freq(fin_freq, "Final Word Cloud")
                        except Exception:
                            st.warning("Could not render final word cloud.")
                        
                        st.caption("Difference (Added Words)")
                        try:
                            render_wordcloud_from_freq(diff_freq, "Added Words Word Cloud")
                        except Exception:
                            st.warning("Could not render difference word cloud.")
                
                st.divider()
                with st.expander("📖 Readability (Flesch Reading Ease)", expanded=False):
                    try:
                        import textstat
                        fre_src = textstat.flesch_reading_ease(state.get('source_text', ''))
                        fre_fin = textstat.flesch_reading_ease(state.get('final_translation', ''))
                        c1, c2 = st.columns(2)
                        c1.metric("Source", f"{fre_src:.1f}")
                        c2.metric("Final", f"{fre_fin:.1f}", delta=f"{fre_fin - fre_src:+.1f}")
                    except Exception:
                        st.info("Install `textstat` for readability: `pip install textstat`")
                
                st.divider()
                if languages_equivalent(state['source_language'], state['target_language']):
                    with st.expander("🎯 BERTScore (same-language refine mode)", expanded=True):
                        if not BERT_AVAILABLE:
                            st.info("Install `bert-score` to enable this metric: `pip install bert-score`")
                        else:
                            refs = state.get('source_text', '').strip()
                            cands = state.get('final_translation', '').strip()
                            if refs and cands:
                                bs = compute_bertscore(candidate=cands, reference=refs)
                                if bs:
                                    try:
                                        fig3, ax3 = plt.subplots()
                                        ax3.barh(["Precision", "Recall", "F1"], [bs["precision"], bs["recall"], bs["f1"]])
                                        ax3.set_xlim(0, 1)
                                        ax3.set_xlabel("Score")
                                        ax3.set_title("BERTScore")
                                        st.pyplot(fig3)
                                    except Exception:
                                        st.warning("Could not render BERTScore chart.")
                                else:
                                    st.warning("BERTScore could not be computed.")
                            else:
                                st.info("Provide both source and final text to compute BERTScore.")
                        
                        if state.get('bertscore_history'):
                            st.divider()
                            st.markdown("#### 🎯 BERTScore Refinement History")
                            
                            history_df = pd.DataFrame(state['bertscore_history'])
                            
                            try:
                                fig_history, ax_history = plt.subplots()
                                ax_history.plot(history_df['attempt'], history_df['f1'], 'o-', label='F1', linewidth=2)
                                ax_history.plot(history_df['attempt'], history_df['precision'], 's-', label='Precision', alpha=0.7)
                                ax_history.plot(history_df['attempt'], history_df['recall'], '^-', label='Recall', alpha=0.7)
                                ax_history.axhline(y=0.8, color='r', linestyle='--', label='Target (0.8)')
                                ax_history.set_xlabel('Refinement Attempt')
                                ax_history.set_ylabel('Score')
                                ax_history.set_title('BERTScore Evolution')
                                ax_history.legend()
                                ax_history.grid(True, alpha=0.3)
                                ax_history.set_ylim(0, 1)
                                st.pyplot(fig_history)
                                
                                st.caption(f"Total attempts: {len(history_df)} | Final F1: {history_df.iloc[-1]['f1']:.3f}")
                            except Exception:
                                st.warning("Could not render BERTScore refinement history chart.")
                else:
                    st.caption("BERTScore is only shown for same-language refine runs.")
                
                st.divider()
                with st.expander("📋 Issues by Stage", expanded=False):
                    issue_counts = {
                        "Literal": len(state.get('literal_issues', [])),
                        "Cultural/Register": len(state.get('cultural_issues', [])),
                        "Tone": len(state.get('tone_issues', [])),
                        "Technical": len(state.get('technical_issues', [])),
                        "Literary": len(state.get('literary_issues', [])),
                    }
                    try:
                        df_issues = pd.DataFrame(
                            {"Stage": list(issue_counts.keys()), "Count": list(issue_counts.values())}
                        ).set_index("Stage")
                        st.bar_chart(df_issues)
                    except Exception:
                        st.warning("Could not render issues chart.")
            else:
                st.info("Visualizations are disabled.")
        
        # History Tab
        with tab_history:
            st.subheader("📚 History")
            if st.session_state.history:
                for i, item in enumerate(reversed(st.session_state.history), 1):
                    with st.expander(f"Run {len(st.session_state.history) - i + 1} – {item['timestamp'][:19]}"):
                        st.write(f"**Model:** {item.get('model','Unknown')}")
                        st.write(f"**Target:** {item.get('target_lang','Unknown')}")
                        st.write(f"**Source Preview:** {item['source']}")
                        if st.button(f"Load this output", key=f"load_{i}"):
                            st.session_state.translation_state = item['result']
                            st.session_state.edited_translation = item['result'].get('final_translation','')
                            st.rerun()
            else:
                st.info("No history yet")
        
        # Entity Tab (only if enabled)
        if st.session_state.enable_entity_tracking and tab_entities:
            with tab_entities:
                st.subheader("🎯 Entity Analysis & Tracking")
                
                if state:
                    entity_tracker = st.session_state.entity_tracker
                    
                    if 'source_entities' not in state or state['source_entities'] is None:
                        state['source_entities'] = entity_tracker.extract_entities(state.get('source_text', ''))
                    if 'translated_entities' not in state or state['translated_entities'] is None:
                        state['translated_entities'] = entity_tracker.extract_entities(state.get('final_translation', ''))
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("### 📖 Source Entities")
                        source_entities = state.get('source_entities', [])
                        
                        if source_entities:
                            st.metric("Total Entities", len(source_entities))
                            
                            entity_types = {}
                            for entity in source_entities:
                                if entity['type'] not in entity_types:
                                    entity_types[entity['type']] = []
                                entity_types[entity['type']].append(entity)
                            
                            for entity_type, entities in entity_types.items():
                                emoji = entity_tracker.entity_types[entity_type]['emoji']
                                with st.expander(f"{emoji} {entity_type.capitalize()} ({len(entities)})"):
                                    for entity in entities[:20]:
                                        badge = "📚" if entity.get('from_glossary') else "🔮"
                                        st.write(f"• **{entity['name']}** {badge} (×{entity['count']})")
                                        if entity.get('description') and entity['description'] != 'Auto-detected':
                                            st.caption(entity['description'])
                        else:
                            st.info("No entities detected in source")
                    
                    with col2:
                        st.markdown("### 📝 Translated Entities")
                        translated_entities = state.get('translated_entities', [])
                        
                        if translated_entities:
                            st.metric("Total Entities", len(translated_entities))
                            
                            entity_types = {}
                            for entity in translated_entities:
                                if entity['type'] not in entity_types:
                                    entity_types[entity['type']] = []
                                entity_types[entity['type']].append(entity)
                            
                            for entity_type, entities in entity_types.items():
                                emoji = entity_tracker.entity_types[entity_type]['emoji']
                                with st.expander(f"{emoji} {entity_type.capitalize()} ({len(entities)})"):
                                    for entity in entities[:20]:
                                        badge = "📚" if entity.get('from_glossary') else "🔮"
                                        st.write(f"• **{entity['name']}** {badge} (×{entity['count']})")
                                        if entity.get('description') and entity['description'] != 'Auto-detected':
                                            st.caption(entity['description'])
                        else:
                            st.info("No entities detected in translation")
                    
                    if source_entities and translated_entities:
                        st.divider()
                        st.markdown("### 📊 Entity Preservation Analysis")
                        
                        source_names = {e['name'].lower() for e in source_entities}
                        translated_names = {e['name'].lower() for e in translated_entities}
                        
                        preserved = source_names & translated_names
                        lost = source_names - translated_names
                        added = translated_names - source_names
                        
                        col1, col2, col3, col4 = st.columns(4)
                        
                        with col1:
                            preservation_rate = len(preserved) / len(source_names) * 100 if source_names else 0
                            st.metric("Preservation Rate", f"{preservation_rate:.1f}%")
                        
                        with col2:
                            st.metric("Preserved", len(preserved))
                        
                        with col3:
                            st.metric("Lost", len(lost))
                        
                        with col4:
                            st.metric("Added", len(added))
                        
                        if lost:
                            with st.expander(f"⚠️ Lost Entities ({len(lost)})", expanded=len(lost) <= 5):
                                for name in list(lost)[:30]:
                                    entity = next((e for e in source_entities if e['name'].lower() == name), None)
                                    if entity:
                                        emoji = entity_tracker.entity_types[entity['type']]['emoji']
                                        st.write(f"{emoji} {entity['name']}")
                        
                        if added:
                            with st.expander(f"➕ Added Entities ({len(added)})"):
                                for name in list(added)[:30]:
                                    entity = next((e for e in translated_entities if e['name'].lower() == name), None)
                                    if entity:
                                        emoji = entity_tracker.entity_types[entity['type']]['emoji']
                                        st.write(f"{emoji} {entity['name']}")
                        
                        st.divider()
                        st.markdown("### 🌐 Entity Network Visualization")
                        
                        viz_choice = st.radio(
                            "Select entities to visualize",
                            ["Source Entities", "Translated Entities", "All Entities"],
                            horizontal=True
                        )
                        
                        if viz_choice == "Source Entities":
                            viz_entities = source_entities[:30]
                        elif viz_choice == "Translated Entities":
                            viz_entities = translated_entities[:30]
                        else:
                            all_entities = {}
                            for e in source_entities + translated_entities:
                                if e['name'] not in all_entities:
                                    all_entities[e['name']] = e
                            viz_entities = list(all_entities.values())[:30]
                        
                        if viz_entities:
                            entity_tracker.visualize_network(viz_entities)
                        
                        if PLOTLY_AVAILABLE:
                            st.divider()
                            st.markdown("### 📈 Entity Frequency Comparison")
                            
                            all_entity_names = {}
                            for e in source_entities:
                                all_entity_names[e['name']] = {'source': e['count'], 'translated': 0, 'type': e['type']}
                            for e in translated_entities:
                                if e['name'] in all_entity_names:
                                    all_entity_names[e['name']]['translated'] = e['count']
                                else:
                                    all_entity_names[e['name']] = {'source': 0, 'translated': e['count'], 'type': e['type']}
                            
                            sorted_entities = sorted(all_entity_names.items(),
                                                   key=lambda x: x[1]['source'] + x[1]['translated'],
                                                   reverse=True)[:20]
                            
                            if sorted_entities:
                                df_data = []
                                for name, counts in sorted_entities:
                                    df_data.append({
                                        'Entity': name,
                                        'Source': counts['source'],
                                        'Translated': counts['translated']
                                    })
                                
                                df = pd.DataFrame(df_data)
                                
                                fig = go.Figure()
                                fig.add_trace(go.Bar(
                                    name='Source',
                                    y=df['Entity'],
                                    x=df['Source'],
                                    orientation='h',
                                    marker_color='#3b82f6'
                                ))
                                fig.add_trace(go.Bar(
                                    name='Translated',
                                    y=df['Entity'],
                                    x=df['Translated'],
                                    orientation='h',
                                    marker_color='#10b981'
                                ))
                                
                                fig.update_layout(
                                    title='Top 20 Entities: Source vs Translated',
                                    xaxis_title='Occurrences',
                                    barmode='group',
                                    height=600,
                                    margin=dict(l=150)
                                )
                                
                                st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("Complete a translation to see entity analysis")

if __name__ == "__main__":
    asyncio.run(main())