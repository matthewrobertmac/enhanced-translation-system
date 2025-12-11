"""
Enhanced Multi-Agent Translation Workflow with LangGraph and LangSmith
A sophisticated translation system with cultural adaptation, literary editing, comprehensive monitoring, and visuals.

Features:
- INTELLIGENT PLANNING AGENT - dynamically selects required agents
- 7 specialized translation agents with distinct roles (including BERTScore validator)
- **SMART SEMANTIC CACHING** - 5-10x speedup on similar content
- **CONFIDENCE SCORES** - Multi-metric translation quality assessment
- **DIFF VISUALIZATION** - Visual comparison between agent versions
- **ALTERNATIVE TRANSLATIONS** - Generate and compare multiple variants
- Support for OpenAI and Anthropic models
- Optional LangSmith tracing and monitoring
- Comprehensive agent feedback system
- File upload support (txt, docx, md)
- Multiple export formats (txt, docx, md)
- Critical passage flagging and review
- Safe same-language (e.g., English→English) refinement mode
- BERTScore validation with iterative refinement
- Visualizations: word counts, sentence-length histograms, readability, issue counts, BERTScore bars
- Word clouds: Source, Final, and Difference (words added)
- Entity tracking and network visualization
"""
from typing import TypedDict, Annotated, List, Dict, Optional, Literal
from langgraph.graph import StateGraph, END
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langchain_core.language_models import BaseChatModel
import operator
from datetime import datetime
import json
import os
import io
import traceback
import re
import collections
import difflib
import asyncio
import hashlib
import pickle

# === New: viz imports ===
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt

# === New: wordcloud (optional) ===
try:
    from wordcloud import WordCloud, STOPWORDS
    WORDCLOUD_AVAILABLE = True
except Exception:
    WORDCLOUD_AVAILABLE = False

# === Entity tracking imports ===
try:
    import networkx as nx
    NETWORKX_AVAILABLE = True
except ImportError:
    NETWORKX_AVAILABLE = False

try:
    import plotly.express as px
    import plotly.graph_objects as go
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False

# === Sentence embeddings for caching ===
try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

# ===== Language Guardrail =====
LANGUAGE_GUARDRAIL = (
    "STRICT LANGUAGE GUARDRAIL:\n"
    "- Reply ONLY in the target language specified.\n"
    "- Do NOT include any words or phrases in other languages.\n"
    "- Any notes, bullets, or headings must also be in the target language.\n"
)

# =====================
# CONFIDENCE SCORER
# =====================
class ConfidenceScorer:
    """Score translation confidence using multiple signals"""
    
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
    
    async def score_translation(
        self, 
        source: str, 
        translation: str, 
        context: dict
    ) -> Dict[str, float]:
        """Generate confidence scores"""
        
        scores = {}
        
        # 1. Semantic similarity (if same-language or embeddings available)
        try:
            from bert_score import score as bert_score_fn
            if languages_equivalent(context.get('source_lang', ''), context.get('target_lang', '')):
                P, R, F1 = bert_score_fn([translation], [source], lang="en", rescale_with_baseline=True)
                scores['semantic_fidelity'] = float(F1.mean().item())
        except:
            scores['semantic_fidelity'] = 0.75  # Default reasonable score
        
        # 2. Length ratio (should be reasonable)
        src_len = len(source.split())
        tgt_len = len(translation.split())
        ratio = tgt_len / src_len if src_len > 0 else 0
        # Ideal ratio is 0.8-1.2
        scores['length_ratio'] = max(0, 1.0 - min(abs(ratio - 1.0), 0.5) * 2)
        
        # 3. Fluency (use language model perplexity proxy)
        prompt = f"""Rate the fluency/naturalness of this {context.get('target_lang', 'target')} text on a scale of 0.0-1.0.

Text: {translation[:500]}

Respond with ONLY a number between 0.0 and 1.0, no explanation."""
        
        try:
            response = await self.llm.ainvoke([HumanMessage(content=prompt)])
            fluency_str = response.content.strip()
            # Extract first number found
            import re
            numbers = re.findall(r'0\.\d+|1\.0|0|1', fluency_str)
            if numbers:
                fluency = float(numbers[0])
                scores['fluency'] = min(max(fluency, 0.0), 1.0)
            else:
                scores['fluency'] = 0.75
        except:
            scores['fluency'] = 0.75  # Default to reasonable score
        
        # 4. Terminology consistency
        entities = context.get('entities', [])
        if entities:
            preserved = sum(1 for e in entities if e['name'].lower() in translation.lower())
            total = len(entities)
            scores['terminology'] = preserved / total if total > 0 else 1.0
        else:
            scores['terminology'] = 0.85  # Default when no entities
        
        # 5. Overall confidence (weighted average)
        weights = {
            'semantic_fidelity': 0.35,
            'length_ratio': 0.15,
            'fluency': 0.35,
            'terminology': 0.15
        }
        
        overall = sum(scores.get(k, 0.5) * v for k, v in weights.items())
        scores['overall'] = overall
        
        return scores

# =====================
# DIFF VISUALIZATION
# =====================
def create_diff_visualization(text1: str, text2: str, title1: str, title2: str) -> str:
    """Visual diff with highlighting"""
    
    # Split into words for better granularity
    words1 = text1.split()
    words2 = text2.split()
    
    d = difflib.Differ()
    diff = list(d.compare(words1, words2))
    
    html_output = f"""
    <div style='background: #f8f9fa; padding: 20px; border-radius: 8px; margin: 10px 0;'>
        <div style='display: grid; grid-template-columns: 1fr 1fr; gap: 20px;'>
            <div>
                <h4 style='color: #dc3545; margin-bottom: 10px;'>← {title1}</h4>
                <div style='background: white; padding: 15px; border-radius: 5px; border-left: 4px solid #dc3545; line-height: 1.8;'>
    """
    
    # Build left side (removals and unchanged)
    for word in diff:
        if word.startswith('- '):
            html_output += f"<span style='background: #ffcdd2; padding: 2px 4px; margin: 1px; text-decoration: line-through;'>{word[2:]}</span> "
        elif word.startswith('  '):
            html_output += f"<span style='color: #666;'>{word[2:]}</span> "
    
    html_output += """
                </div>
            </div>
            <div>
                <h4 style='color: #28a745; margin-bottom: 10px;'>{} →</h4>
                <div style='background: white; padding: 15px; border-radius: 5px; border-left: 4px solid #28a745; line-height: 1.8;'>
    """.format(title2)
    
    # Build right side (additions and unchanged)
    for word in diff:
        if word.startswith('+ '):
            html_output += f"<span style='background: #c8e6c9; padding: 2px 4px; margin: 1px; font-weight: 500;'>{word[2:]}</span> "
        elif word.startswith('  '):
            html_output += f"<span style='color: #666;'>{word[2:]}</span> "
    
    html_output += """
                </div>
            </div>
        </div>
    </div>
    """
    
    return html_output

def calculate_change_rate(text1: str, text2: str) -> float:
    """Calculate percentage of text that changed"""
    similarity = difflib.SequenceMatcher(None, text1, text2).ratio()
    return (1 - similarity) * 100

# =====================
# ALTERNATIVE TRANSLATION GENERATOR
# =====================
class AlternativeTranslationGenerator:
    """Generate alternative translation variants"""
    
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
    
    async def generate_alternatives(
        self,
        state: "TranslationState",  # Use string for forward reference
        num_alternatives: int = 3
    ) -> List[Dict]:
        """Generate alternative translations with different strategies"""
        
        alternatives = []
        
        # Strategy 1: More conservative (lower temperature)
        alternatives.append(await self._generate_variant(
            state,
            temperature=0.2,
            strategy="conservative",
            description="More literal and conservative"
        ))
        
        # Strategy 2: Balanced (medium temperature)
        alternatives.append(await self._generate_variant(
            state,
            temperature=0.5,
            strategy="balanced",
            description="Balanced between literal and natural"
        ))
        
        # Strategy 3: More creative (higher temperature)
        if num_alternatives >= 3:
            alternatives.append(await self._generate_variant(
                state,
                temperature=0.8,
                strategy="creative",
                description="More natural and idiomatic"
            ))
        
        return alternatives
    
    async def _generate_variant(
        self,
        state: Dict,  # Use Dict instead of TranslationState
        temperature: float,
        strategy: str,
        description: str
    ) -> Dict:
        """Generate a single translation variant"""
        
        try:
            # Create LLM with specific temperature
            if hasattr(self.llm, 'model_name'):
                from langchain_openai import ChatOpenAI
                variant_llm = ChatOpenAI(
                    model=self.llm.model_name,
                    temperature=temperature,
                    api_key=self.llm.openai_api_key,
                    timeout=120
                )
            else:
                from langchain_anthropic import ChatAnthropic
                variant_llm = ChatAnthropic(
                    model=self.llm.model,
                    temperature=temperature,
                    anthropic_api_key=self.llm.anthropic_api_key,
                    timeout=120
                )
            
            same_lang = languages_equivalent(state['source_language'], state['target_language'])
            
            if same_lang:
                prompt = f"""Refine this {state['target_language']} text. Strategy: {strategy}.

{LANGUAGE_GUARDRAIL}

Source text:
{state['source_text']}

Provide ONLY the refined text, no explanation."""
            else:
                prompt = f"""Translate from {state['source_language']} to {state['target_language']}. Strategy: {strategy}.

{LANGUAGE_GUARDRAIL}

Source text:
{state['source_text']}

Target audience: {state['target_audience']}
Genre: {state.get('genre', 'General')}

Provide ONLY the translation, no explanation."""
            
            response = await variant_llm.ainvoke([HumanMessage(content=prompt)])
            translation = response.content.strip()
            
            return {
                'strategy': strategy,
                'description': description,
                'temperature': temperature,
                'translation': translation,
                'error': False
            }
        
        except Exception as e:
            return {
                'strategy': strategy,
                'description': description,
                'temperature': temperature,
                'translation': f"Error generating alternative: {str(e)}",
                'error': True
            }
        
# =====================
# SEMANTIC TRANSLATION CACHE
# =====================
class SemanticTranslationCache:
    """Cache translations with semantic similarity-based retrieval"""
    
    def __init__(self, cache_dir: str = ".translation_cache"):
        self.cache_dir = cache_dir
        self.cache = {}  # Full translation cache
        self.sentence_cache = {}  # Sentence-level cache
        self.embeddings = {}  # Store embeddings for similarity search
        self.similarity_threshold = 0.85
        self.embedding_model = None
        self.stats = {
            'hits': 0,
            'misses': 0,
            'partial_hits': 0,
            'similar_hits': 0
        }
        
        # Create cache directory
        os.makedirs(cache_dir, exist_ok=True)
        
        # Load existing cache
        self.load_cache()
    
    def get_embedding_model(self):
        """Lazy load embedding model"""
        if self.embedding_model is None and SENTENCE_TRANSFORMERS_AVAILABLE:
            self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        return self.embedding_model
    
    def get_embedding(self, text: str) -> Optional[np.ndarray]:
        """Get or compute embedding for text"""
        if not SENTENCE_TRANSFORMERS_AVAILABLE:
            return None
        
        text_hash = self.hash_text(text)
        
        if text_hash not in self.embeddings:
            model = self.get_embedding_model()
            if model:
                self.embeddings[text_hash] = model.encode(text)
        
        return self.embeddings.get(text_hash)
    
    def hash_text(self, text: str) -> str:
        """Create hash for text"""
        return hashlib.md5(text.encode()).hexdigest()
    
    def make_cache_key(self, text: str, context: dict) -> str:
        """Create cache key from text and context"""
        context_str = f"{context.get('source_lang')}_{context.get('target_lang')}_{context.get('audience')}_{context.get('genre')}"
        text_hash = self.hash_text(text)
        return f"{text_hash}_{hashlib.md5(context_str.encode()).hexdigest()}"
    
    def context_matches(self, cached_context: dict, query_context: dict) -> bool:
        """Check if contexts are compatible"""
        return (
            cached_context.get('source_lang') == query_context.get('source_lang') and
            cached_context.get('target_lang') == query_context.get('target_lang') and
            cached_context.get('audience') == query_context.get('audience') and
            cached_context.get('genre') == query_context.get('genre')
        )
    
    async def get_cached_translation(
        self, 
        text: str, 
        context: dict
    ) -> Optional[Dict]:
        """Check cache for translation"""
        cache_key = self.make_cache_key(text, context)
        
        # 1. EXACT MATCH (fastest - O(1))
        if cache_key in self.cache:
            self.stats['hits'] += 1
            cached = self.cache[cache_key]
            return {
                'type': 'exact',
                'translation': cached['translation'],
                'metadata': cached['metadata'],
                'speedup': '100x'
            }
        
        # 2. SIMILAR MATCH (still fast - using embeddings)
        if SENTENCE_TRANSFORMERS_AVAILABLE:
            similar = await self.find_similar_translation(text, context)
            if similar:
                self.stats['similar_hits'] += 1
                return {
                    'type': 'similar',
                    'translation': similar['translation'],
                    'similarity': similar['similarity'],
                    'original_text': similar['original_text'],
                    'speedup': '50x',
                    'needs_refinement': True
                }
        
        # 3. SENTENCE-LEVEL CACHE (partial speedup)
        sentence_hits = await self.get_cached_sentences(text, context)
        if sentence_hits and len(sentence_hits['cached']) > len(sentence_hits['all']) * 0.3:
            self.stats['partial_hits'] += 1
            return {
                'type': 'partial',
                'cached_sentences': sentence_hits['cached'],
                'uncached_sentences': sentence_hits['uncached'],
                'speedup': f"{int(len(sentence_hits['cached']) / len(sentence_hits['all']) * 100)}%"
            }
        
        # 4. CACHE MISS
        self.stats['misses'] += 1
        return None
    
    async def find_similar_translation(
        self, 
        text: str, 
        context: dict
    ) -> Optional[Dict]:
        """Find semantically similar cached translation"""
        if not SENTENCE_TRANSFORMERS_AVAILABLE:
            return None
        
        query_embedding = self.get_embedding(text)
        if query_embedding is None:
            return None
        
        best_match = None
        best_similarity = 0.0
        
        for cache_key, cached_data in self.cache.items():
            # Check context compatibility
            if not self.context_matches(cached_data['context'], context):
                continue
            
            # Get cached text embedding
            cached_text = cached_data['source_text']
            cached_embedding = self.get_embedding(cached_text)
            
            if cached_embedding is not None:
                # Compute cosine similarity
                similarity = float(np.dot(query_embedding, cached_embedding) / 
                                 (np.linalg.norm(query_embedding) * np.linalg.norm(cached_embedding)))
                
                if similarity > self.similarity_threshold and similarity > best_similarity:
                    best_similarity = similarity
                    best_match = {
                        'translation': cached_data['translation'],
                        'similarity': similarity,
                        'original_text': cached_text
                    }
        
        return best_match
    
    async def get_cached_sentences(
        self, 
        text: str, 
        context: dict
    ) -> Optional[Dict]:
        """Get cached translations for individual sentences"""
        sentences = [s.strip() + '.' for s in text.split('.') if s.strip()]
        
        cached_sentences = []
        uncached_sentences = []
        
        for i, sentence in enumerate(sentences):
            sent_key = self.make_cache_key(sentence, context)
            
            if sent_key in self.sentence_cache:
                cached_sentences.append({
                    'index': i,
                    'original': sentence,
                    'translation': self.sentence_cache[sent_key]['translation']
                })
            else:
                uncached_sentences.append({
                    'index': i,
                    'original': sentence
                })
        
        if not cached_sentences:
            return None
        
        return {
            'all': sentences,
            'cached': cached_sentences,
            'uncached': uncached_sentences
        }
    
    def store_translation(
        self, 
        text: str, 
        translation: str, 
        context: dict,
        metadata: Optional[dict] = None
    ):
        """Store translation in cache"""
        cache_key = self.make_cache_key(text, context)
        
        # Store full translation
        self.cache[cache_key] = {
            'source_text': text,
            'translation': translation,
            'context': context,
            'metadata': metadata or {},
            'timestamp': datetime.now().isoformat()
        }
        
        # Store sentence-level cache
        sentences = [s.strip() + '.' for s in text.split('.') if s.strip()]
        translated_sentences = [s.strip() + '.' for s in translation.split('.') if s.strip()]
        
        if len(sentences) == len(translated_sentences):
            for src_sent, tgt_sent in zip(sentences, translated_sentences):
                sent_key = self.make_cache_key(src_sent, context)
                self.sentence_cache[sent_key] = {
                    'translation': tgt_sent,
                    'timestamp': datetime.now().isoformat()
                }
        
        # Compute and store embedding
        if SENTENCE_TRANSFORMERS_AVAILABLE:
            self.get_embedding(text)  # This will cache it
        
        # Auto-save periodically
        if len(self.cache) % 10 == 0:
            self.save_cache()
    
    def save_cache(self):
        """Save cache to disk"""
        try:
            # Save main cache
            cache_file = os.path.join(self.cache_dir, 'translations.pkl')
            with open(cache_file, 'wb') as f:
                pickle.dump(self.cache, f)
            
            # Save sentence cache
            sent_cache_file = os.path.join(self.cache_dir, 'sentences.pkl')
            with open(sent_cache_file, 'wb') as f:
                pickle.dump(self.sentence_cache, f)
            
            # Save embeddings
            if self.embeddings:
                emb_file = os.path.join(self.cache_dir, 'embeddings.pkl')
                with open(emb_file, 'wb') as f:
                    pickle.dump(self.embeddings, f)
            
            # Save stats
            stats_file = os.path.join(self.cache_dir, 'stats.json')
            with open(stats_file, 'w') as f:
                json.dump(self.stats, f, indent=2)
            
            return True
        except Exception as e:
            print(f"Cache save failed: {str(e)}")
            return False
    
    def load_cache(self):
        """Load cache from disk"""
        try:
            # Load main cache
            cache_file = os.path.join(self.cache_dir, 'translations.pkl')
            if os.path.exists(cache_file):
                with open(cache_file, 'rb') as f:
                    self.cache = pickle.load(f)
            
            # Load sentence cache
            sent_cache_file = os.path.join(self.cache_dir, 'sentences.pkl')
            if os.path.exists(sent_cache_file):
                with open(sent_cache_file, 'rb') as f:
                    self.sentence_cache = pickle.load(f)
            
            # Load embeddings
            emb_file = os.path.join(self.cache_dir, 'embeddings.pkl')
            if os.path.exists(emb_file):
                with open(emb_file, 'rb') as f:
                    self.embeddings = pickle.load(f)
            
            # Load stats
            stats_file = os.path.join(self.cache_dir, 'stats.json')
            if os.path.exists(stats_file):
                with open(stats_file, 'r') as f:
                    self.stats = json.load(f)
            
            return True
        except Exception as e:
            print(f"Cache load failed: {str(e)}")
            return False
    
    def clear_cache(self):
        """Clear all caches"""
        self.cache = {}
        self.sentence_cache = {}
        self.embeddings = {}
        self.stats = {
            'hits': 0,
            'misses': 0,
            'partial_hits': 0,
            'similar_hits': 0
        }
        
        # Delete cache files
        try:
            for filename in ['translations.pkl', 'sentences.pkl', 'embeddings.pkl', 'stats.json']:
                filepath = os.path.join(self.cache_dir, filename)
                if os.path.exists(filepath):
                    os.remove(filepath)
        except Exception as e:
            print(f"Cache clear failed: {str(e)}")
    
    def get_stats(self) -> dict:
        """Get cache statistics"""
        total_requests = sum(self.stats.values())
        hit_rate = (self.stats['hits'] + self.stats['similar_hits'] + self.stats['partial_hits']) / max(total_requests, 1) * 100
        
        return {
            **self.stats,
            'total_entries': len(self.cache),
            'sentence_entries': len(self.sentence_cache),
            'embedding_entries': len(self.embeddings),
            'hit_rate': hit_rate,
            'total_requests': total_requests
        }
    
    def export_cache(self) -> dict:
        """Export cache for sharing"""
        return {
            'cache': self.cache,
            'sentence_cache': self.sentence_cache,
            'stats': self.stats,
            'exported_at': datetime.now().isoformat()
        }
    
    def import_cache(self, data: dict):
        """Import cache from export"""
        self.cache.update(data.get('cache', {}))
        self.sentence_cache.update(data.get('sentence_cache', {}))
        self.save_cache()

# =====================
# ENTITY TRACKER CLASS
# =====================
import csv
from collections import Counter, defaultdict

class EntitiesTracker:
    """Entity tracking and visualization system with network graphs"""
    
    def __init__(self):
        self.entity_types = {
            'person': {'emoji': '👤', 'color': '#3b82f6'},
            'location': {'emoji': '📍', 'color': '#10b981'},
            'organization': {'emoji': '🏢', 'color': '#f59e0b'},
            'date': {'emoji': '📅', 'color': '#8b5cf6'},
            'custom': {'emoji': '🏷️', 'color': '#ef4444'}
        }
        
        # Backend storage 
        self.entity_glossary = {}
        self.extracted_entities = []

    
    def extract_entities(self, text: str) -> List[Dict]:
        """Extract entities from text using glossary and NER patterns"""
        if not text:
            return []
        
        entities = []
        
        # Extract from glossary
        for term, info in self.entity_glossary.items():
            pattern = re.compile(r'\b' + re.escape(term) + r'\b', re.IGNORECASE)
            matches = pattern.findall(text)
            
            # Check aliases
            alias_count = 0
            for alias in info.get('aliases', []):
                alias_pattern = re.compile(r'\b' + re.escape(alias) + r'\b', re.IGNORECASE)
                alias_matches = alias_pattern.findall(text)
                alias_count += len(alias_matches)
            
            total_count = len(matches) + alias_count
            
            if total_count > 0:
                entities.append({
                    'name': term,
                    'type': info.get('type', 'custom'),
                    'count': total_count,
                    'description': info.get('description', ''),
                    'from_glossary': True
                })
        
        # Auto-detection patterns
        # Person names
        person_pattern = r'\b([A-Z][a-z]+ [A-Z][a-z]+)\b'
        for match in re.finditer(person_pattern, text):
            name = match.group(1)
            if name not in [e['name'] for e in entities]:
                count = len(re.findall(r'\b' + re.escape(name) + r'\b', text))
                entities.append({
                    'name': name,
                    'type': 'person',
                    'count': count,
                    'description': 'Auto-detected',
                    'auto_detected': True
                })
        
        # Organizations
        org_pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Inc|Corp|LLC|Ltd|Company|Group|Foundation))\b'
        for match in re.finditer(org_pattern, text):
            org = match.group(1)
            if org not in [e['name'] for e in entities]:
                count = len(re.findall(r'\b' + re.escape(org) + r'\b', text))
                entities.append({
                    'name': org,
                    'type': 'organization',
                    'count': count,
                    'description': 'Auto-detected',
                    'auto_detected': True
                })
        
        return entities
    
    def upload_glossary(self, file) -> bool:
        """Process uploaded glossary file"""
        try:
            content = file.read()
            
            if file.name.endswith('.json'):
                new_glossary = json.loads(content)
            elif file.name.endswith('.csv'):
                csv_data = io.StringIO(content.decode('utf-8'))
                reader = csv.DictReader(csv_data)
                new_glossary = {}
                for row in reader:
                    term = row.get('term', '').strip()
                    if term:
                        new_glossary[term] = {
                            'type': row.get('type', 'custom'),
                            'description': row.get('description', ''),
                            'aliases': [a.strip() for a in row.get('aliases', '').split(',') if a.strip()]
                        }
            else:
                lines = content.decode('utf-8').split('\n')
                new_glossary = {}
                for line in lines:
                    term = line.strip()
                    if term:
                        new_glossary[term] = {'type': 'custom', 'description': '', 'aliases': []}
            self.entity_glossary.update(new_glossary)
            return True
        except Exception as e:
            print(f"Error processing glossary: {str(e)}")
            return False
def visualize_network(self, entities: List[Dict]):
    """Create network visualization and return a plotly Figure object."""
    if not NETWORKX_AVAILABLE or not PLOTLY_AVAILABLE:
        print("Install networkx and plotly for network visualization")
        return None
    
    # Create graph
    G = nx.Graph()
    
    # Add nodes
    for entity in entities:
        G.add_node(entity['name'],
                   type=entity['type'],
                   count=entity['count'],
                   description=entity.get('description', ''))
    
    # Add edges (simplified co-occurrence)
    for i, e1 in enumerate(entities):
        for e2 in entities[i+1:]:
            weight = min(e1['count'], e2['count'])
            G.add_edge(e1['name'], e2['name'], weight=weight)
    
    # Layout
    pos = nx.spring_layout(G, k=2, iterations=50)
    
    # Edge traces
    edge_traces = []
    for edge in G.edges(data=True):
        x0, y0 = pos[edge[0]]
        x1, y1 = pos[edge[1]]
        trace = go.Scatter(
            x=[x0, x1, None],
            y=[y0, y1, None],
            mode='lines',
            line=dict(width=0.5, color='#888'),
            hoverinfo='none'
        )
        edge_traces.append(trace)
    
    # Node traces
    node_x, node_y, node_text, node_color, node_size = [], [], [], [], []
        
    for node in G.nodes():
        x, y = pos[node]
        node_x.append(x)
        node_y.append(y)
        
        entity = next(e for e in entities if e['name'] == node)
        node_text.append(f"{node}<br>Type: {entity['type']}<br>Count: {entity['count']}")
        node_color.append(self.entity_types[entity['type']]['color'])
        node_size.append(10 + min(entity['count'] * 3, 50))
    
    node_trace = go.Scatter(
        x=node_x, y=node_y,
        mode='markers+text',
        text=[n for n in G.nodes()],
        textposition="top center",
        hoverinfo='text',
        hovertext=node_text,
        marker=dict(
            size=node_size,
            color=node_color,
            line=dict(width=2, color='white')
        )
    )
    
    fig = go.Figure(
        data=edge_traces + [node_trace],
        layout=go.Layout(
            title='Entity Network',
            showlegend=False,
            hovermode='closest',
            margin=dict(b=0, l=0, r=0, t=40),
            xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
            yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
            height=600
        )
    )
    
    return fig

# ---- Helpers for same-language paths ----
def normalize_lang_label(label: str) -> str:
    """Normalize a language label like 'English (US)' -> 'english'."""
    core = label.split("(")[0].strip().lower()
    return core

def languages_equivalent(src: str, tgt: str) -> bool:
    return normalize_lang_label(src) == normalize_lang_label(tgt)

def mode_phrase(src: str, tgt: str) -> str:
    """Human-readable short descriptor for UI headers."""
    if languages_equivalent(src, tgt):
        return f"{tgt} – refine"
    return f"{src} → {tgt}"

def split_notes(content: str, labels: List[str]) -> (str, Optional[str], Optional[str]):
    """
    Split content on the first matching label in labels list (e.g., ["TRANSLATOR NOTES:", "EDITOR NOTES:"]).
    Returns (main_text, found_label, notes_text).
    """
    for lab in labels:
        if lab in content:
            parts = content.split(lab, 1)
            return parts[0].strip(), lab, parts[1].strip()
    return content.strip(), None, None

# LangSmith integration (optional)
try:
    from langsmith import Client
    LANGSMITH_AVAILABLE = True
except ImportError:
    LANGSMITH_AVAILABLE = False

# Model imports
try:
    from langchain_openai import ChatOpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

try:
    from langchain_anthropic import ChatAnthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

# Document processing
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# BERTScore (optional; only used in same-language mode)
try:
    from bert_score import score as bert_score_fn
    BERT_AVAILABLE = True
except Exception:
    BERT_AVAILABLE = False

# Checkpointer
from langgraph.checkpoint.memory import MemorySaver

# Caching
from langchain_core.caches import InMemoryCache
from langchain_core.globals import set_llm_cache
set_llm_cache(InMemoryCache())

# =====================
# AGENT ROLE DESCRIPTIONS
# =====================
AGENT_DESCRIPTIONS = {
    "planning": {
        "name": "Workflow Planning Specialist",
        "emoji": "📋",
        "role": "Intelligent Agent Selection & Routing",
        "description": """
        **Primary Responsibility:** Analyze source text and dynamically determine which agents are needed.
        **Key Functions:**
        - Assess text complexity and characteristics
        - Detect technical, cultural, and literary elements
        - Select optimal agent sequence for the task
        - Estimate time and cost savings
        - Provide reasoning for all decisions
        """,
        "expertise": ["Text analysis", "Workflow optimization", "Resource efficiency", "Quality assurance"]
    },
    "literal_translator": {
        "name": "Baseline Specialist",
        "emoji": "🔤",
        "role": "Foundational Accuracy (or faithful baseline in refine mode)",
        "description": """
        **Primary Responsibility:** Provides the baseline pass:
        - Translation mode: create an accurate word-for-word draft.
        - Refine mode (same-language): perform a faithful copy-edit that preserves meaning and structure.
        **Key Functions:**
        - Preserve semantic content and important structure
        - Identify idioms, ambiguities, cultural expressions
        - Preserve technical terms and proper nouns
        - Document challenges for downstream agents
        """,
        "expertise": ["Lexical mapping", "Semantic preservation", "Grammatical structure", "Technical terminology"]
    },
    "cultural_adapter": {
        "name": "Cultural Localization Expert",
        "emoji": "🌍",
        "role": "Cross-Cultural Bridge (or register-localization in refine mode)",
        "description": """
        **Primary Responsibility:** Make content culturally natural in target context.
        - In refine mode, ensure register and references suit the target variety (e.g., US vs UK English).
        """,
        "expertise": ["Localization", "Idiomatic expressions", "Cross-cultural communication", "Audience adaptation"]
    },
    "tone_specialist": {
        "name": "Tone & Voice Consistency Director",
        "emoji": "🎭",
        "role": "Stylistic Harmony & Readability",
        "description": "Ensure consistent tone, voice, pacing, and readability.",
        "expertise": ["Stylistics", "Rhetorical analysis", "Reader psychology", "Narrative voice"]
    },
    "technical_reviewer": {
        "name": "Technical Accuracy Auditor",
        "emoji": "🔬",
        "role": "Precision & Factual Integrity",
        "description": "Verify notation, terminology, measurements, and formats.",
        "expertise": ["Technical writing", "Scientific notation", "Terminology", "Quality assurance"]
    },
    "literary_editor": {
        "name": "Literary Style & Excellence Editor",
        "emoji": "✍️",
        "role": "Publication-Ready Prose Crafting",
        "description": "Elevate prose to publication quality without altering meaning.",
        "expertise": ["Literary criticism", "Creative writing", "Stylistics", "Publishing standards"]
    },
    "quality_controller": {
        "name": "Master Quality Synthesizer",
        "emoji": "✅",
        "role": "Final Integration & Excellence Assurance",
        "description": "Integrate all contributions and approve the final version.",
        "expertise": ["Editorial judgment", "Synthesis", "Holistic QA"]
    },
    "bertscore_validator": {
        "name": "Semantic Fidelity Validator",
        "emoji": "🎯",
        "role": "Semantic Similarity Assurance (same-language only)",
        "description": """
        **Primary Responsibility:** Ensure semantic fidelity in same-language refinement.
        - Only active when source and target languages are equivalent
        - Validates BERTScore ≥ 0.8 for semantic preservation
        - Iteratively refines output to meet threshold
        **Key Functions:**
        - Compute BERTScore (P/R/F1) against source
        - Identify semantic drift areas
        - Apply targeted refinements without over-editing
        - Preserve meaning while improving clarity
        """,
        "expertise": ["Semantic analysis", "Embedding comparison", "Iterative refinement", "Fidelity validation"]
    }
}

# =====================
# State Definition (UPDATED WITH NEW FIELDS)
# =====================
class TranslationState(TypedDict):
    source_text: str
    source_language: str
    target_language: str
    target_audience: str
    genre: str
    # Versions
    literal_translation: str
    cultural_adaptation: str
    tone_adjustment: str
    technical_review_version: str
    literary_polish: str
    final_translation: str
    # Issues
    literal_issues: List[Dict]
    cultural_issues: List[Dict]
    tone_issues: List[Dict]
    technical_issues: List[Dict]
    literary_issues: List[Dict]
    # Critical passages
    critical_passages: List[Dict]
    # Workflow tracking
    agent_notes: Annotated[List[str], operator.add]
    agent_decisions: List[Dict]
    human_feedback: Optional[str]
    revision_count: int
    needs_human_review: bool
    # Entity tracking (Optional)
    source_entities: Optional[List[Dict]]
    translated_entities: Optional[List[Dict]]
    entity_preservation_rate: Optional[float]
    # BERTScore tracking
    bertscore_attempts: int
    bertscore_history: List[Dict]
    # Planning fields
    agent_plan: List[str]
    agent_plan_reasoning: Dict[str, str]
    planning_analysis: Dict[str, any]
    skipped_agents: List[str]
    estimated_complexity: str
    current_agent_index: int
    planning_enabled: bool
    # Cache tracking
    cache_hit: Optional[str]
    cache_speedup: Optional[str]
    # NEW: Confidence scores
    confidence_scores: Optional[Dict[str, float]]
    # Metadata
    started_at: str
    completed_at: Optional[str]

# =====================
# File Processing Utilities
# =====================
def read_uploaded_file(uploaded_file) -> str:
    """Read content from uploaded file"""
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        if file_extension in ('txt', 'md'):
            return uploaded_file.read().decode('utf-8')
        elif file_extension == 'docx':
            if not DOCX_AVAILABLE:
                print("python-docx not installed. Install with: pip install python-docx")
                return ""
            doc = Document(uploaded_file)
            return '\n\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            print(f"Unsupported file type: {file_extension}")
            return ""
    except Exception as e:
        print(f"Error reading file: {str(e)}")
        return ""

def create_docx_file(text: str, title: str = "Translation") -> io.BytesIO:
    """Create a formatted Word document"""
    if not DOCX_AVAILABLE:
        print("python-docx not installed. Install with: pip install python-docx")
        return None
    
    doc = Document()
    title_paragraph = doc.add_heading(title, level=1)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    timestamp = doc.add_paragraph()
    timestamp.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    for para in text.split('\n\n'):
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def create_markdown_file(text: str, title: str = "Translation") -> str:
    return f"""# {title}

**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

---

{text}
"""

def extract_critical_passages(text: str, issues_list: List[Dict]) -> List[Dict]:
    """Extract critical passages that need review based on agent feedback"""
    critical = []
    for issue in issues_list:
        content = issue.get('content', '')
        issue_type = issue.get('type', '')
        agent = issue.get('agent', '')
        
        if any(k in content.lower() for k in ['critical', 'warning', 'error', 'ambiguous', 'unclear', 'needs review']):
            sentences = text.split('.')
            for i, sentence in enumerate(sentences):
                if len(sentence.strip()) > 10:
                    critical.append({
                        'passage': sentence.strip() + '.',
                        'issue': content[:200],
                        'agent': agent,
                        'type': issue_type,
                        'sentence_index': i
                    })
    
    return critical[:10]

# =====================
# Optional: Language Reinforcement
# =====================
async def reinforce_language(llm: BaseChatModel, text: str, target_language: str) -> str:
    """Lightweight fixer to ensure output remains in the target language only."""
    try:
        resp = await llm.ainvoke([
            SystemMessage(content="Ensure the following text is in the specified language only, without code-switching."),
            HumanMessage(content=f"{LANGUAGE_GUARDRAIL}\n\nTarget language: {target_language}\n\nText:\n{text}\n\nReturn the corrected text in {target_language}, with no extra commentary.")
        ])
        return resp.content.strip()
    except Exception:
        return text

# =====================
# BERTScore Utility (same-language only)
# =====================
def compute_bertscore(candidate: str, reference: str) -> Optional[Dict[str, float]]:
    """Compute BERTScore (P/R/F1) if package is available; returns None otherwise."""
    if not BERT_AVAILABLE:
        return None
    try:
        P, R, F1 = bert_score_fn([candidate], [reference], lang="en", rescale_with_baseline=True)
        return {
            "precision": float(P.mean().item()),
            "recall": float(R.mean().item()),
            "f1": float(F1.mean().item()),
        }
    except Exception:
        return None

# =====================
# Visualization helpers
# =====================
def sentence_lengths(text: str) -> List[int]:
    """Return per-sentence word counts for a text."""
    if not text:
        return []
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    return [len(s.split()) for s in sentences if s.strip()]

def nonzero_bins(max_len: int) -> List[int]:
    """Nice 5-word bins for histograms."""
    step = 5
    upper = max(5, ((max_len + step - 1) // step) * step)
    return list(range(0, upper + step, step))

_WORD_RE = re.compile(r"[A-Za-zÀ-ÖØ-öø-ÿ']+")

def tokenize_words(text: str) -> List[str]:
    if not text:
        return []
    return [w.lower() for w in _WORD_RE.findall(text)]

def word_frequencies(text: str, stopwords: Optional[set] = None) -> Dict[str, int]:
    tokens = tokenize_words(text)
    if stopwords is None:
        stopwords = set()
    cnt = collections.Counter(w for w in tokens if w not in stopwords and len(w) > 1)
    return dict(cnt)

def render_wordcloud_from_freq(freqs: Dict[str, int], title: str):
    if not freqs:
        print(f"No tokens to display for **{title}**.")
        return
    
    wc = WordCloud(
        width=1000,
        height=500,
        background_color="white",
        collocations=False
    ).generate_from_frequencies(freqs)
    
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.imshow(wc, interpolation="bilinear")
    ax.axis("off")
    ax.set_title(title)
    return fig

def text_similarity(a: str, b: str) -> float:
    return difflib.SequenceMatcher(None, a or "", b or "").ratio()

# =====================
# PLANNING AGENT
# =====================
class PlanningAgent:
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
        self.name = "Workflow Planning Specialist"
        self.emoji = "📋"
    
    async def analyze_and_plan(self, state: TranslationState) -> TranslationState:
        """Analyze text and create execution plan"""
        if state.get('agent_plan'):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Plan already exists - skipping")
            return state
        
        attempt = 0
        max_attempts = 2
        
        while attempt < max_attempts:
            try:
                # Check if planning is disabled (manual override)
                if not state.get('planning_enabled', True):
                    state['agent_plan'] = [
                        'literal_translator', 'cultural_adapter', 'tone_specialist',
                        'technical_reviewer', 'literary_editor', 'finalize', 'bertscore_validator'
                    ]
                    state['agent_plan_reasoning'] = {
                        'literal_translator': '✅ Required - baseline',
                        'cultural_adapter': '✅ Included - planning disabled',
                        'tone_specialist': '✅ Included - planning disabled',
                        'technical_reviewer': '✅ Included - planning disabled',
                        'literary_editor': '✅ Included - planning disabled',
                        'finalize': '✅ Required - final review',
                        'bertscore_validator': '✅ Included - planning disabled'
                    }
                    state['skipped_agents'] = []
                    state['estimated_complexity'] = 'full'
                    state['current_agent_index'] = 0
                    state['planning_analysis'] = {'mode': 'manual_override'}
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Planning disabled - running all agents")
                    break
                
                source_text = state['source_text']
                source_lang = state['source_language']
                target_lang = state['target_language']
                audience = state['target_audience']
                genre = state.get('genre', 'General')
                same_lang = languages_equivalent(source_lang, target_lang)
                
                # Check for manual overrides from session state
                force_cultural = False
                force_tone = False
                force_technical = False
                force_literary = False
                force_bertscore = False                
                system_prompt = f"""You are an expert translation workflow planner. Analyze the source text and determine which translation agents are needed.

AVAILABLE AGENTS:
1. literal_translator (ALWAYS REQUIRED) - Creates initial translation/refinement
2. cultural_adapter - Handles idioms, cultural references, localization
3. tone_specialist - Ensures voice consistency and readability
4. technical_reviewer - Validates technical accuracy, terminology, units
5. literary_editor - Elevates prose quality for publication
6. finalize (ALWAYS REQUIRED) - Final quality synthesis
7. bertscore_validator - Semantic fidelity validation (same-language only)

ANALYSIS CRITERIA:
- Text length: {len(source_text)} characters, ~{len(source_text.split())} words
- Language pair: {source_lang} → {target_lang} (same-language: {same_lang})
- Target audience: {audience}
- Genre: {genre}
- Look for: technical terms, cultural references, idioms, literary elements, tone consistency issues

SELECTION GUIDELINES:
- literal_translator & finalize: ALWAYS include
- cultural_adapter: Include if cross-language OR idioms/cultural content OR same-lang with register differences
- tone_specialist: Include if >500 words OR inconsistent tone detected OR multiple sections
- technical_reviewer: Include if technical terminology, formulas, units, citations present
- literary_editor: Include if literary genre OR narrative elements OR high-polish audience
- bertscore_validator: ALWAYS include if same-language mode. DO NOT SKIP.

EFFICIENCY GOAL: Include only agents that will meaningfully improve output. When uncertain, include the agent (favor quality over speed).

OUTPUT FORMAT: Return ONLY valid JSON (no markdown, no backticks):
{{
    "required_agents": ["literal_translator", "cultural_adapter", "finalize"],
    "reasoning": {{
        "literal_translator": "✅ Required - baseline",
        "cultural_adapter": "✅ Detected 2 idioms and cross-language pair",
        "tone_specialist": "⏭️ Skipped - short text (120 words), consistent tone",
        "technical_reviewer": "⏭️ Skipped - no technical content",
        "literary_editor": "⏭️ Skipped - casual genre",
        "finalize": "✅ Required - final review",
        "bertscore_validator": "⏭️ Skipped - cross-language"
    }},
    "analysis": {{
        "length": 120,
        "complexity_score": 2.3,
        "technical_terms_count": 0,
        "cultural_references_count": 2,
        "tone_consistency": "high",
        "literary_elements": "minimal"
    }},
    "estimated_complexity": "simple"
}}
"""
                
                user_prompt = f"""Analyze this text and create an optimal translation workflow plan:

SOURCE TEXT:
{source_text[:2000]}{'...' if len(source_text) > 2000 else ''}

Return ONLY the JSON plan, no other text."""
                
                response = await self.llm.ainvoke([
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=user_prompt)
                ])
                
                content = response.content.strip()
                content = re.sub(r'```json\s*', '', content)
                content = re.sub(r'```\s*$', '', content)
                
                plan_data = json.loads(content)
                
                required_agents = plan_data.get('required_agents', [])
                reasoning = plan_data.get('reasoning', {})
                analysis = plan_data.get('analysis', {})
                complexity = plan_data.get('estimated_complexity', 'moderate')
                
                # Apply manual overrides
                if force_cultural and 'cultural_adapter' not in required_agents:
                    required_agents.insert(1, 'cultural_adapter')
                    reasoning['cultural_adapter'] = '✅ Forced by user override'
                
                if force_tone and 'tone_specialist' not in required_agents:
                    insert_pos = 2 if 'cultural_adapter' in required_agents else 1
                    required_agents.insert(insert_pos, 'tone_specialist')
                    reasoning['tone_specialist'] = '✅ Forced by user override'
                
                if force_technical and 'technical_reviewer' not in required_agents:
                    finalize_pos = required_agents.index('finalize') if 'finalize' in required_agents else len(required_agents)
                    required_agents.insert(finalize_pos, 'technical_reviewer')
                    reasoning['technical_reviewer'] = '✅ Forced by user override'
                
                if force_literary and 'literary_editor' not in required_agents:
                    finalize_pos = required_agents.index('finalize') if 'finalize' in required_agents else len(required_agents)
                    required_agents.insert(finalize_pos, 'literary_editor')
                    reasoning['literary_editor'] = '✅ Forced by user override'
                
                if force_bertscore and 'bertscore_validator' not in required_agents and same_lang:
                    required_agents.append('bertscore_validator')
                    reasoning['bertscore_validator'] = '✅ Forced by user override'
                
                # Ensure required agents are present
                if 'literal_translator' not in required_agents:
                    required_agents.insert(0, 'literal_translator')
                if 'finalize' not in required_agents:
                    required_agents.append('finalize')
                
                # Enforce bertscore_validator always runs in same-language mode
                if same_lang and BERT_AVAILABLE and 'bertscore_validator' not in required_agents:
                    required_agents.append('bertscore_validator')
                    reasoning['bertscore_validator'] = '✅ Always required in same-language mode'
                
                all_agents = ['literal_translator', 'cultural_adapter', 'tone_specialist',
                             'technical_reviewer', 'literary_editor', 'finalize', 'bertscore_validator']
                skipped = [a for a in all_agents if a not in required_agents]
                
                state['agent_plan'] = required_agents
                state['agent_plan_reasoning'] = reasoning
                state['planning_analysis'] = analysis
                state['skipped_agents'] = skipped
                state['estimated_complexity'] = complexity
                state['current_agent_index'] = 0
                
                total_agents = len(all_agents)
                used_agents = len(required_agents)
                savings_pct = int((1 - used_agents / total_agents) * 100)
                
                state['agent_notes'].append(
                    f"{self.emoji} {self.name}: Plan created - {used_agents}/{total_agents} agents "
                    f"({savings_pct}% savings) - Complexity: {complexity}"
                )
                break
            
            except Exception as e:
                attempt += 1
                if attempt < max_attempts:
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Retry attempt {attempt} after error: {str(e)}")
                else:
                    print(f"Planning agent failed after {max_attempts} attempts, using conservative defaults: {str(e)}")
                    state['agent_plan'] = [
                        'literal_translator', 'cultural_adapter', 'tone_specialist',
                        'technical_reviewer', 'literary_editor', 'finalize'
                    ]
                    if same_lang and BERT_AVAILABLE:
                        state['agent_plan'].append('bertscore_validator')
                    
                    state['agent_plan_reasoning'] = {
                        a: '✅ Included (fallback mode)' for a in state['agent_plan']
                    }
                    state['skipped_agents'] = []
                    state['estimated_complexity'] = 'unknown'
                    state['current_agent_index'] = 0
                    state['planning_analysis'] = {'error': str(e), 'mode': 'fallback'}
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Planning failed - using all agents (safe mode)")
        
        return state

# =====================
# FIXED ROUTER FUNCTION
# =====================
def route_to_next_agent(state: TranslationState) -> str:
    """Determine next agent based on the plan"""
    agent_plan = state.get('agent_plan', [])
    current_index = state.get('current_agent_index', 0)
    
    if current_index >= len(agent_plan):
        return END
    
    next_agent = agent_plan[current_index]
    
    return next_agent

# =====================
# AGENT CLASSES - Each increments index
# =====================
class LiteralTranslationAgent:
    def __init__(self, llm: BaseChatModel, cache: SemanticTranslationCache):
        self.llm = llm
        self.cache = cache
        self.name = "Baseline Specialist"
        self.emoji = "🔤"
    
    async def translate(self, state: TranslationState) -> TranslationState:
        if state.get('literal_translation'):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Already complete - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        # Check cache first
        context = {
            'source_lang': state['source_language'],
            'target_lang': state['target_language'],
            'audience': state['target_audience'],
            'genre': state.get('genre', 'General')
        }
        
        cached = await self.cache.get_cached_translation(state['source_text'], context)
        
        if cached and cached['type'] == 'exact':
            state['literal_translation'] = cached['translation']
            state['literal_issues'] = []
            state['cache_hit'] = 'exact'
            state['cache_speedup'] = cached['speedup']
            state['agent_notes'].append(f"⚡ {self.emoji} {self.name}: CACHE HIT (exact match, {cached['speedup']} faster)")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        if cached and cached['type'] == 'similar':
            state['agent_notes'].append(f"⚡ {self.emoji} {self.name}: CACHE HIT (similar match, {cached['similarity']:.2f} similarity)")
            # Use similar translation as starting point
            base_translation = cached['translation']
            state['cache_hit'] = 'similar'
            state['cache_speedup'] = cached['speedup']
        else:
            base_translation = None
        
        attempt = 0
        max_attempts = 2
        
        while attempt < max_attempts:
            try:
                same_lang = languages_equivalent(state['source_language'], state['target_language'])
                system_prompt = (
                    "You provide the baseline pass:\n"
                    "- Translation mode: accurate, faithful literal translation.\n"
                    "- Refine mode: faithful copy-edit that preserves meaning and structure.\n\n"
                    + LANGUAGE_GUARDRAIL
                )
                
                if same_lang:
                    user_prompt = f"""Refine the following text in {state['target_language']} without changing its meaning.

{LANGUAGE_GUARDRAIL}

**GOAL (REFINE MODE):**
- Preserve semantics, improve clarity and microstructure
- Flag idioms/ambiguities for downstream agents
- Keep terminology and proper nouns intact
- OUTPUT FORMAT: Provide the refined text first, then add "EDITOR NOTES:" (notes must be in {state['target_language']}).

**TEXT TO REFINE:**
{state['source_text']}

**AUDIENCE**: {state['target_audience']}
**GENRE**: {state.get('genre', 'General')}
"""
                    if self.enable_entity_awareness:
                        if state.get('source_entities'):
                            entity_list = "\n".join([
                            f"- {e['name']} ({e['type']})"
                            for e in state['source_entities'][:15]
                            ])
                        user_prompt += (
                            f"\n\n**IMPORTANT ENTITIES TO PRESERVE:**\n{entity_list}\n"
                         )

                    if self.enable_entity_awareness:
                        if state.get('source_entities'):
                            entity_list = "\n".join([f"- {e['name']} ({e['type']})" for e in state['source_entities'][:15]])
                            user_prompt += f"\n\n**IMPORTANT ENTITIES TO PRESERVE:**\n{entity_list}\n"
                else:
                    user_prompt = f"""Translate the following text from {state['source_language']} to {state['target_language']}.

{LANGUAGE_GUARDRAIL}

**CRITICAL INSTRUCTIONS:**
1. Translate with maximum fidelity to the original meaning.
2. Maintain sentence structure initially.
3. Flag idioms/ambiguities for downstream agents.
4. OUTPUT FORMAT: Provide the literal translation first, then "TRANSLATOR NOTES:" (notes must be in {state['target_language']}).

**SOURCE TEXT:**
{state['source_text']}

**TARGET AUDIENCE**: {state['target_audience']}
**GENRE**: {state.get('genre', 'General')}
"""
# Add entity-awareness instructions (if enabled)
                    if self.enable_entity_awareness:
                        if state.get('source_entities'):
                            entity_list = "\n".join([
                            f"- {e['name']} ({e['type']})"
                            for e in state['source_entities'][:15]
                            ])
                            user_prompt += (
                            f"\n\n**IMPORTANT ENTITIES TO PRESERVE:**\n{entity_list}\n"
                            )

                    if self.enable_entity_awareness:
                        if state.get('source_entities'):
                            entity_list = "\n".join([f"- {e['name']} ({e['type']})" for e in state['source_entities'][:15]])
                            user_prompt += f"\n\n**IMPORTANT ENTITIES TO PRESERVE:**\n{entity_list}\n"
                
                response = await self.llm.ainvoke([SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)])
                content = response.content
                
                translation, found_label, notes = split_notes(content, ["TRANSLATOR NOTES:", "EDITOR NOTES:"])
                
                issues = []
                if notes:
                    issues.append({
                        "agent": self.name,
                        "type": (found_label[:-1].lower().replace(" ", "_") if found_label else "notes"),
                        "content": notes
                    })
                
                translation = await reinforce_language(self.llm, translation, state['target_language'])
                
                state['literal_translation'] = translation
                state['literal_issues'] = issues
                
                # Store in cache
                self.cache.store_translation(
                    state['source_text'],
                    translation,
                    context,
                    {'agent': self.name, 'timestamp': datetime.now().isoformat()}
                )
                
                state['agent_notes'].append(f"{self.emoji} {self.name}: Baseline {'refinement' if same_lang else 'literal translation'} complete")
                break
            
            except Exception as e:
                attempt += 1
                if attempt < max_attempts:
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Retry attempt {attempt} after error: {str(e)}")
                else:
                    print(f"Error in Literal step after {max_attempts} attempts: {str(e)}")
                    state['literal_translation'] = state['source_text']
                    state['literal_issues'] = [{"agent": self.name, "type": "error", "content": str(e)}]
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Error - using source text")
        
        # INCREMENT INDEX
        state['current_agent_index'] = state.get('current_agent_index', 0) + 1
        return state

class CulturalAdaptationAgent:
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
        self.name = "Cultural Localization Expert"
        self.emoji = "🌍"
    
    async def adapt(self, state: TranslationState) -> TranslationState:
        if state.get('cultural_adaptation'):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Already complete - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        previous_text = state.get('literal_translation') or ""
        
        if text_similarity(previous_text, state.get('cultural_adaptation', '')) > 0.95:
            state['agent_notes'].append(f"{self.emoji} {self.name}: Minimal changes detected - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        attempt = 0
        max_attempts = 2
        
        while attempt < max_attempts:
            try:
                same_lang = languages_equivalent(state['source_language'], state['target_language'])
                system_prompt = "You adapt content for cultural/register naturalness in the target context.\n\n" + LANGUAGE_GUARDRAIL
                
                if same_lang:
                    task_text = (
                        "Refine for the target variety's norms (e.g., spelling, idioms, punctuation, register). "
                        "Replace region-specific expressions with appropriate target-variety equivalents."
                    )
                    notes_label = "EDITOR NOTES:"
                else:
                    task_text = (
                        "Replace source-culture idioms with target equivalents, adapt references, "
                        "and adjust communication style for the target audience."
                    )
                    notes_label = "CULTURAL NOTES:"
                
                user_prompt = f"""Work on the following text.

{LANGUAGE_GUARDRAIL}

**TARGET CONTEXT**: {state['target_language']} / Audience: {state['target_audience']}

**TEXT:**
{state['literal_translation']}

**YOUR TASKS:**
- {task_text}

**OUTPUT**: Provide the adapted text, then {notes_label} (notes must be in {state['target_language']})."""
                
                response = await self.llm.ainvoke([SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)])
                content = response.content
                
                adapted, found_label, notes = split_notes(content, ["CULTURAL NOTES:", "EDITOR NOTES:"])
                
                issues = []
                if notes:
                    issues.append({
                        "agent": self.name,
                        "type": (found_label[:-1].lower().replace(" ", "_") if found_label else "notes"),
                        "content": notes
                    })
                
                adapted = await reinforce_language(self.llm, adapted, state['target_language'])
                
                state['cultural_adaptation'] = adapted
                state['cultural_issues'] = issues
                state['agent_notes'].append(f"{self.emoji} {self.name}: {'Register' if same_lang else 'Cultural'} adaptation complete")
                break
            
            except Exception as e:
                attempt += 1
                if attempt < max_attempts:
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Retry attempt {attempt} after error: {str(e)}")
                else:
                    print(f"Error in Cultural step after {max_attempts} attempts: {str(e)}")
                    state['cultural_adaptation'] = state['literal_translation']
                    state['cultural_issues'] = [{"agent": self.name, "type": "error", "content": str(e)}]
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Error - using previous version")
        
        # INCREMENT INDEX
        state['current_agent_index'] = state.get('current_agent_index', 0) + 1
        return state

class ToneConsistencyAgent:
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
        self.name = "Tone & Voice Consistency Director"
        self.emoji = "🎭"
    
    async def adjust_tone(self, state: TranslationState) -> TranslationState:
        if state.get('tone_adjustment'):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Already complete - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        previous_text = state.get('cultural_adaptation') or state.get('literal_translation') or ""
        
        if text_similarity(previous_text, state.get('tone_adjustment', '')) > 0.95:
            state['agent_notes'].append(f"{self.emoji} {self.name}: Minimal changes detected - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        attempt = 0
        max_attempts = 2
        
        while attempt < max_attempts:
            try:
                system_prompt = "You ensure tone/voice consistency and natural readability.\n\n" + LANGUAGE_GUARDRAIL
                
                user_prompt = f"""Adjust this text for tone consistency and readability.

{LANGUAGE_GUARDRAIL}

**AUDIENCE**: {state['target_audience']}

**TEXT:**
{previous_text}

**YOUR TASKS:**
1. Vary sentence length for rhythm
2. Match formality to audience
3. Ensure consistent voice
4. Optimize readability

**OUTPUT**: Provide the adjusted text, then "TONE NOTES:" (notes must be in {state['target_language']})."""
                
                response = await self.llm.ainvoke([SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)])
                content = response.content
                
                adjusted, found_label, notes = split_notes(content, ["TONE NOTES:"])
                
                issues = []
                if notes:
                    issues.append({"agent": self.name, "type": "tone_notes", "content": notes})
                
                adjusted = await reinforce_language(self.llm, adjusted, state['target_language'])
                
                state['tone_adjustment'] = adjusted
                state['tone_issues'] = issues
                state['agent_notes'].append(f"{self.emoji} {self.name}: Tone/readability optimized")
                break
            
            except Exception as e:
                attempt += 1
                if attempt < max_attempts:
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Retry attempt {attempt} after error: {str(e)}")
                else:
                    print(f"Error in Tone step after {max_attempts} attempts: {str(e)}")
                    state['tone_adjustment'] = previous_text
                    state['tone_issues'] = [{"agent": self.name, "type": "error", "content": str(e)}]
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Error - using previous version")
        
        # INCREMENT INDEX
        state['current_agent_index'] = state.get('current_agent_index', 0) + 1
        return state

class TechnicalReviewAgent:
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
        self.name = "Technical Accuracy Auditor"
        self.emoji = "🔬"
    
    async def review(self, state: TranslationState) -> TranslationState:
        if state.get('technical_review_version'):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Already complete - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        previous_text = (state.get('tone_adjustment') or
                         state.get('cultural_adaptation') or
                         state.get('literal_translation') or "")
        
        if text_similarity(previous_text, state.get('technical_review_version', '')) > 0.95:
            state['agent_notes'].append(f"{self.emoji} {self.name}: Minimal changes detected - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        attempt = 0
        max_attempts = 2
        
        while attempt < max_attempts:
            try:
                system_prompt = "You verify notation, terminology, units, and formatting.\n\n" + LANGUAGE_GUARDRAIL
                
                user_prompt = f"""Review the following for technical accuracy.

{LANGUAGE_GUARDRAIL}

**TEXT:**
{previous_text}

**YOUR TASKS:**
1. Verify notation/symbols
2. Check terminology
3. Validate measurements/units
4. Ensure number/date/time formatting

**OUTPUT**: Provide the reviewed text, then "TECHNICAL NOTES:" if corrections made (notes must be in {state['target_language']})."""
                
                response = await self.llm.ainvoke([SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)])
                content = response.content
                
                reviewed, found_label, notes = split_notes(content, ["TECHNICAL NOTES:"])
                
                issues = []
                if notes:
                    issues.append({"agent": self.name, "type": "technical_notes", "content": notes})
                
                needs_review = "NEEDS_REVIEW" in content or "NEEDS REVIEW" in content
                
                reviewed = await reinforce_language(self.llm, reviewed, state['target_language'])
                
                state['technical_review_version'] = reviewed
                state['technical_issues'] = issues
                state['needs_human_review'] = needs_review
                state['agent_notes'].append(f"{self.emoji} {self.name}: Technical review completed")
                break
            
            
            except Exception as e:
                attempt += 1
                if attempt < max_attempts:
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Retry attempt {attempt} after error: {str(e)}")
                else:
                    print(f"Error in Technical step after {max_attempts} attempts: {str(e)}")
                    state['technical_review_version'] = previous_text
                    state['technical_issues'] = [{"agent": self.name, "type": "error", "content": str(e)}]
                    state['needs_human_review'] = False
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Error - using previous version")
        
        # INCREMENT INDEX
        state['current_agent_index'] = state.get('current_agent_index', 0) + 1
        return state

class LiteraryEditorAgent:
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
        self.name = "Literary Style & Excellence Editor"
        self.emoji = "✍️"
    
    async def polish(self, state: TranslationState) -> TranslationState:
        if state.get('literary_polish'):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Already complete - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        previous_text = (state.get('technical_review_version') or
                         state.get('tone_adjustment') or
                         state.get('cultural_adaptation') or
                         state.get('literal_translation') or "")
        
        if text_similarity(previous_text, state.get('literary_polish', '')) > 0.95:
            state['agent_notes'].append(f"{self.emoji} {self.name}: Minimal changes detected - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        attempt = 0
        max_attempts = 2
        
        while attempt < max_attempts:
            try:
                system_prompt = "You elevate prose to publication quality without altering meaning.\n\n" + LANGUAGE_GUARDRAIL
                
                user_prompt = f"""Polish this text to publication quality.

{LANGUAGE_GUARDRAIL}

**TEXT:**
{previous_text}

**AUDIENCE**: {state['target_audience']}

**YOUR TASKS:**
- Eliminate awkward phrasing
- Enhance word choice
- Optimize prose rhythm
- Strengthen imagery
- Maintain meaning

**OUTPUT**: Provide the polished text, then "LITERARY NOTES:" (notes must be in {state['target_language']})."""
                
                response = await self.llm.ainvoke([SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)])
                content = response.content
                
                polished, found_label, notes = split_notes(content, ["LITERARY NOTES:"])
                
                issues = []
                if notes:
                    issues.append({"agent": self.name, "type": "literary_notes", "content": notes})
                
                polished = await reinforce_language(self.llm, polished, state['target_language'])
                
                state['literary_polish'] = polished
                state['literary_issues'] = issues
                state['agent_notes'].append(f"{self.emoji} {self.name}: Literary polish completed")
                break
            
            except Exception as e:
                attempt += 1
                if attempt < max_attempts:
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Retry attempt {attempt} after error: {str(e)}")
                else:
                    print(f"Error in Literary step after {max_attempts} attempts: {str(e)}")
                    state['literary_polish'] = previous_text
                    state['literary_issues'] = [{"agent": self.name, "type": "error", "content": str(e)}]
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Error - using previous version")
        
        # INCREMENT INDEX
        state['current_agent_index'] = state.get('current_agent_index', 0) + 1
        return state

class QualityControlAgent:
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
        self.name = "Master Quality Synthesizer"
        self.emoji = "✅"
    
    async def finalize(self, state: TranslationState) -> TranslationState:
        if state.get('final_translation'):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Already complete - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        previous_text = (state.get('literary_polish') or
                         state.get('technical_review_version') or
                         state.get('tone_adjustment') or
                         state.get('cultural_adaptation') or
                         state.get('literal_translation') or "")
        
        if text_similarity(previous_text, state.get('final_translation', '')) > 0.95:
            state['agent_notes'].append(f"{self.emoji} {self.name}: Minimal changes detected - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        attempt = 0
        max_attempts = 2
        
        while attempt < max_attempts:
            try:
                system_prompt = "You integrate all contributions and output the final publication-ready text.\n\n" + LANGUAGE_GUARDRAIL
                
                all_issues = (
                    state.get('literal_issues', []) +
                    state.get('cultural_issues', []) +
                    state.get('tone_issues', []) +
                    state.get('technical_issues', []) +
                    state.get('literary_issues', [])
                )
                
                latest_version = (
                    state.get('literary_polish') or
                    state.get('technical_review_version') or
                    state.get('tone_adjustment') or
                    state.get('cultural_adaptation') or
                    state.get('literal_translation')
                )
                
                user_prompt = f"""Produce the final, publication-ready text.

{LANGUAGE_GUARDRAIL}

**LATEST VERSION:**
{latest_version}

**OUTPUT**: Provide ONLY the final text without meta-commentary, entirely in {state['target_language']}.
"""
                
                response = await self.llm.ainvoke([SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)])
                content = await reinforce_language(self.llm, response.content.strip(), state['target_language'])
                
                critical_passages = extract_critical_passages(content, all_issues)
                
                state['final_translation'] = content
                state['completed_at'] = datetime.now().isoformat()
                state['critical_passages'] = critical_passages
                state['agent_notes'].append(f"{self.emoji} {self.name}: Final output approved")
                
                # NEW: Calculate confidence scores
                confidence_scorer = ConfidenceScorer(self.llm)
                context = {
                    'source_lang': state['source_language'],
                    'target_lang': state['target_language'],
                    'entities': state.get('source_entities', [])
                }
                confidence = await confidence_scorer.score_translation(
                    state['source_text'],
                    content,
                    context
                )
                state['confidence_scores'] = confidence
                
                if self.enable_entity_tracking:
                    entity_tracker = self.entity_tracker
                    state['translated_entities'] = entity_tracker.extract_entities(content)
                    
                    if state.get('source_entities'):
                        source_names = {e['name'].lower() for e in state['source_entities']}
                        translated_names = {e['name'].lower() for e in state['translated_entities']}
                        if source_names:
                            state['entity_preservation_rate'] = len(source_names & translated_names) / len(source_names)
                
                break
            
            except Exception as e:
                attempt += 1
                if attempt < max_attempts:
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Retry attempt {attempt} after error: {str(e)}")
                else:
                    print(f"Error in Finalize step after {max_attempts} attempts: {str(e)}")
                    state['final_translation'] = latest_version
                    state['completed_at'] = datetime.now().isoformat()
                    state['critical_passages'] = []
                    state['agent_notes'].append(f"{self.emoji} {self.name}: Error - using latest version")
        
        # INCREMENT INDEX
        state['current_agent_index'] = state.get('current_agent_index', 0) + 1
        return state

class BERTScoreValidatorAgent:
    def __init__(self, llm: BaseChatModel):
        self.llm = llm
        self.name = "Semantic Fidelity Validator"
        self.emoji = "🎯"
        self.target_score = 0.8
    
    async def validate_and_refine(self, state: TranslationState) -> TranslationState:
        """Validate BERTScore and refine if needed (same-language mode only)"""
        if state.get('bertscore_history') and state['bertscore_history'][-1]['f1'] >= self.target_score:
            state['agent_notes'].append(f"{self.emoji} {self.name}: Already validated - skipping")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        if not languages_equivalent(state['source_language'], state['target_language']):
            state['agent_notes'].append(f"{self.emoji} {self.name}: Skipped (cross-language mode)")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        if not BERT_AVAILABLE:
            state['agent_notes'].append(f"{self.emoji} {self.name}: Skipped (bert-score not installed)")
            state['current_agent_index'] = state.get('current_agent_index', 0) + 1
            return state
        
        source_text = state['source_text']
        current_text = state['final_translation'] or ""
        
        if 'bertscore_attempts' not in state:
            state['bertscore_attempts'] = 0
        if 'bertscore_history' not in state:
            state['bertscore_history'] = []
        
        attempt = 0
        max_refinements = 20  # Safety cap to prevent infinite loops
        
        while attempt < max_refinements:
            state['bertscore_attempts'] += 1
            attempt += 1
            
            scores = compute_bertscore(current_text, source_text)
            
            if scores is None:
                state['agent_notes'].append(f"{self.emoji} {self.name}: BERTScore computation failed")
                break
            
            f1_score = scores['f1']
            state['bertscore_history'].append({
                'attempt': state['bertscore_attempts'],
                'f1': f1_score,
                'precision': scores['precision'],
                'recall': scores['recall']
            })
            
            if f1_score >= self.target_score:
                state['agent_notes'].append(
                    f"{self.emoji} {self.name}: ✅ BERTScore validated (F1={f1_score:.3f}) after {attempt} attempt(s)"
                )
                state['final_translation'] = current_text
                break
            
            if text_similarity(current_text, source_text) > 0.98:  # Near identical, no point refining
                state['agent_notes'].append(f"{self.emoji} {self.name}: Near-identical to source - accepting (F1={f1_score:.3f})")
                break
            
            state['agent_notes'].append(
                f"{self.emoji} {self.name}: 🔄 Refining (F1={f1_score:.3f} < {self.target_score}, attempt {attempt})"
            )
            
            system_prompt = (
                "You are a semantic fidelity specialist. Your ONLY goal is to increase semantic similarity "
                "to the source text while maintaining natural language quality.\n\n"
                + LANGUAGE_GUARDRAIL
            )
            
            user_prompt = f"""The current refined text has insufficient semantic similarity to the source (BERTScore F1: {f1_score:.3f}, target: {self.target_score}).

{LANGUAGE_GUARDRAIL}

**SOURCE TEXT (preserve its meaning):**
{source_text}

**CURRENT REFINED TEXT (needs closer alignment):**
{current_text}

**YOUR TASK:**
1. Identify where meaning has drifted from the source
2. Adjust ONLY those areas to restore semantic fidelity
3. DO NOT over-edit - preserve what is already good
4. Maintain natural, fluent language
5. Keep technical terms and proper nouns identical

**CRITICAL:** The goal is semantic similarity, not word-for-word copying. Preserve the SOURCE's meaning using natural language.

**OUTPUT:** Provide ONLY the refined text in {state['target_language']}, with no meta-commentary."""
            
            try:
                response = await self.llm.ainvoke([
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=user_prompt)
                ])
                current_text = await reinforce_language(self.llm, response.content.strip(), state['target_language'])
                
            except Exception as e:
                print(f"Error in BERTScore refinement: {str(e)}")
                state['agent_notes'].append(f"{self.emoji} {self.name}: Error during refinement - {str(e)}")
                break
        
        state['final_translation'] = current_text
        
        # INCREMENT INDEX
        state['current_agent_index'] = state.get('current_agent_index', 0) + 1
        return state

# =====================
# LLM Initialization Helper
# =====================
def initialize_llm(
    provider: Literal["openai", "anthropic"],
    model: str,
    api_key: str,
    temperature: float = 0.3
) -> BaseChatModel:
    if provider == "openai":
        if not OPENAI_AVAILABLE:
            raise ImportError("OpenAI not available. Install with: pip install langchain-openai")
        return ChatOpenAI(model=model, temperature=temperature, api_key=api_key, timeout=120)
    elif provider == "anthropic":
        if not ANTHROPIC_AVAILABLE:
            raise ImportError("Anthropic not available. Install with: pip install langchain-anthropic")
        return ChatAnthropic(model=model, temperature=temperature, api_key=api_key, timeout=120)
    else:
        raise ValueError(f"Unsupported provider: {provider}")

# =====================
# LangSmith Setup
# =====================
def setup_langsmith(api_key: Optional[str], project_name: str = "translation-pipeline"):
    if not api_key:
        return False
    if not LANGSMITH_AVAILABLE:
        print("LangSmith not installed. Install with: pip install langsmith")
        return False
    
    try:
        os.environ["LANGCHAIN_TRACING_V2"] = "true"
        os.environ["LANGCHAIN_API_KEY"] = api_key
        os.environ["LANGCHAIN_PROJECT"] = project_name
        Client(api_key=api_key)
        return True
    except Exception as e:
        print(f"LangSmith setup failed: {str(e)}")
        return False

# =====================
# Graph Construction WITH DYNAMIC ROUTING
# =====================
def build_translation_graph(llm: BaseChatModel, cache: SemanticTranslationCache) -> StateGraph:
    planning_agent = PlanningAgent(llm)
    literal_agent = LiteralTranslationAgent(llm, cache)
    cultural_agent = CulturalAdaptationAgent(llm)
    tone_agent = ToneConsistencyAgent(llm)
    technical_agent = TechnicalReviewAgent(llm)
    literary_agent = LiteraryEditorAgent(llm)
    qc_agent = QualityControlAgent(llm)
    bertscore_agent = BERTScoreValidatorAgent(llm)
    
class TranslationPipeline:    
    def __init__(self, llm: BaseChatModel):
        # Core model
        self.llm = llm

        # Backend flags
        self.enable_entity_awareness = False
        self.enable_entity_tracking = False

        # Backend-managed objects
        self.entity_tracker = EntitiesTracker()
        self.semantic_cache = SemanticTranslationCache()

    def build_workflow(self):
        """Builds and compiles the full translation pipeline workflow."""
        
        # Instantiate agents
        planning_agent = PlanningAgent(self.llm)
        literal_agent = LiteralTranslationAgent(self.llm, self.semantic_cache)
        cultural_agent = CulturalAdaptationAgent(self.llm)
        tone_agent = ToneConsistencyAgent(self.llm)
        technical_agent = TechnicalReviewAgent(self.llm)
        literary_agent = LiteraryEditorAgent(self.llm)
        qc_agent = QualityControlAgent(self.llm)
        bertscore_agent = BERTScoreValidatorAgent(self.llm)

        # Define workflow
        workflow = StateGraph(TranslationState)

        # Add nodes
        workflow.add_node("planning", planning_agent.analyze_and_plan)
        workflow.add_node("literal_translator", literal_agent.translate)
        workflow.add_node("cultural_adapter", cultural_agent.adapt)
        workflow.add_node("tone_specialist", tone_agent.adjust_tone)
        workflow.add_node("technical_reviewer", technical_agent.review)
        workflow.add_node("literary_editor", literary_agent.polish)
        workflow.add_node("finalize", qc_agent.finalize)
        workflow.add_node("bertscore_validator", bertscore_agent.validate_and_refine)

        # Entry node
        workflow.set_entry_point("planning")

        # Route transitions
        workflow.add_conditional_edges("planning", route_to_next_agent)
        workflow.add_conditional_edges("literal_translator", route_to_next_agent)
        workflow.add_conditional_edges("cultural_adapter", route_to_next_agent)
        workflow.add_conditional_edges("tone_specialist", route_to_next_agent)
        workflow.add_conditional_edges("technical_reviewer", route_to_next_agent)
        workflow.add_conditional_edges("literary_editor", route_to_next_agent)
        workflow.add_conditional_edges("finalize", route_to_next_agent)
        workflow.add_conditional_edges("bertscore_validator", route_to_next_agent)

        checkpointer = MemorySaver()
        return workflow.compile(checkpointer=checkpointer)

    async def run(self, state: TranslationState):
        """Execute full translation pipeline."""
        workflow = self.build_workflow()
        return await workflow.ainvoke(
            state,
            config={
                "thread_id": state.get("thread_id", "default-thread"),
                "checkpoint_ns": "translation",
            }
        )

